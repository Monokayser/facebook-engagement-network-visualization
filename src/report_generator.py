"""Generate the academic Markdown, DOCX, and PDF reports from verified results."""

from __future__ import annotations

import json
from collections.abc import Iterable
from copy import copy
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab import rl_config
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.config import ANALYSIS_SUMMARY, REPORT, STATIC, STUDENT

NAVY = RGBColor(23, 50, 77)
BLUE = RGBColor(57, 115, 168)
GRAY = RGBColor(82, 96, 109)
LIGHT = "E9EEF3"
rl_config.invariant = True


def _fmt(value: float, digits: int = 2) -> str:
    return f"{value:,.{digits}f}"


def _load_summary(path: Path = ANALYSIS_SUMMARY) -> dict[str, Any]:
    with path.open(encoding="utf-8") as stream:
        return json.load(stream)


def _transport_matrix_rows(transport: dict[str, Any]) -> list[list[Any]]:
    """Reconstruct the weighted matrix from verified summary edge records."""

    nodes = transport["node_order"]
    values = {(node, node): 0 for node in nodes}
    for edge in transport["weighted_edges"]:
        source = edge["source"]
        target = edge["target"]
        distance = edge["distance_km"]
        values[(source, target)] = distance
        values[(target, source)] = distance
    return [
        [node, *[values.get((node, other), 0) for other in nodes]] for node in nodes
    ]


def _exercise_markdown(summary: dict[str, Any]) -> str:
    """Create detailed academic sections for all seven exercises."""

    exercises = summary["exercises"]
    transport = exercises["exercise_1"]
    ppi = exercises["exercise_2"]
    bipartite = exercises["exercise_3"]
    weighted = exercises["exercise_4"]
    models = exercises["exercise_5"]
    social = exercises["exercise_6"]
    domain = exercises["exercise_7"]

    matrix_header = "| City | " + " | ".join(transport["node_order"]) + " |"
    matrix_rule = "|---|" + "|".join("---:" for _ in transport["node_order"]) + "|"
    matrix_rows = "\n".join(
        "| " + " | ".join(str(value) for value in row) + " |"
        for row in _transport_matrix_rows(transport)
    )
    course_rows = "\n".join(
        f"| {row['course']} | {row['enrollments']} |"
        for row in bipartite["course_degrees"]
    )
    weight_rows = "\n".join(
        f"| {weight} | {count} |"
        for weight, count in weighted["weight_frequencies"].items()
    )
    model_rows = "\n".join(
        (
            f"| {row['Model']} | {row['Nodes']} | {row['Edges']} | "
            f"{row['Density']:.3f} | {row['Average Degree']:.3f} | "
            f"{row['Average Clustering']:.3f} | "
            f"{row['Average Shortest Path']:.3f} | "
            f"{row['Connected Components']} | {row['Maximum Degree']} | "
            f"{row['Degree Standard Deviation']:.3f} |"
        )
        for row in models["records"]
    )
    group_rows = "\n".join(
        f"| {group} | {count} |"
        for group, count in social["interest_group_counts"].items()
    )
    centrality_rows = "\n".join(
        (
            f"| {row['label']} | {row['node_type']} | {row['degree']} | "
            f"{row['betweenness_centrality']:.4f} | "
            f"{row['closeness_centrality']:.4f} | {row['pagerank']:.4f} |"
        )
        for row in domain["top_centrality_records"]
    )
    top_users = ", ".join(
        f"{row['user']} ({row['in_degree']})"
        for row in social["highest_in_degree_users"]
    )

    return f"""### 8.1 Exercise 1: Weighted adjacency representations

#### 8.1.1 Objective and analytical question

The relationship between an adjacency list and a weighted adjacency matrix was examined through the supplied `G_transport` network. Particular attention was given to whether the matrix was symmetric and to what that symmetry implied about the direction of the modeled routes.

#### 8.1.2 Graph construction and stored attributes

Seven Bangladeshi cities were represented by nodes, and nine teaching routes were represented by undirected edges. A distance in kilometers was stored in the `weight` attribute of every edge. The graph contains {transport['nodes']} nodes, {transport['edges']} edges, and a density of {transport['density']:.4f}. Dhaka was connected to {transport['highest_degree']} other cities and therefore received the largest degree.

#### 8.1.3 Analytical procedure

A labeled adjacency list was produced from NetworkX neighbor iterators. A weighted matrix was generated with `nx.to_pandas_adjacency`, using the documented node order and edge `weight` attribute. Symmetry was tested by comparing the matrix with its transpose through `numpy.allclose`. A zero off the diagonal was interpreted as the absence of a recorded direct route.

#### 8.1.4 Verified results

**Table 2. Weighted adjacency matrix for `G_transport` (kilometers)**

{matrix_header}
{matrix_rule}
{matrix_rows}

The symmetry test returned **{transport['matrix_symmetric']}**. The {transport['longest_route']['source']}-{transport['longest_route']['target']} edge was the longest at {transport['longest_route']['distance_km']} km, while the {transport['shortest_route']['source']}-{transport['shortest_route']['target']} edge was the shortest at {transport['shortest_route']['distance_km']} km. The sum of the nine unique edge distances was {transport['total_recorded_distance_km']:,} km.

#### 8.1.5 Academic interpretation

Symmetry was produced because every route was encoded as an undirected edge with one shared weight. The distance from city u to city v was therefore written into both matrix cells (u,v) and (v,u). An asymmetric matrix would be expected under one-way travel or direction-dependent cost. The adjacency list was better suited to neighbor inspection, while the matrix was better suited to exact pairwise lookup and numerical validation.

#### 8.1.6 Limitations and verification evidence

The graph was treated as a classroom illustration rather than a complete transportation model. Road conditions, route alternatives, travel time, and geographic validation were outside its scope. Reproduction evidence was saved in `outputs/tables/g_transport_adjacency_matrix.csv`, and matrix symmetry is covered by an automated test.

### 8.2 Exercise 2: Six-layout comparison

#### 8.2.1 Objective and controlled comparison

Spring, circular, shell, spectral, Kamada-Kawai, and random layouts were compared for the same `G_ppi` topology. Node identity, size, color, edge set, labels, and figure scale were held constant so that only placement was changed.

#### 8.2.2 Graph construction

The graph was reproduced with `nx.watts_strogatz_graph(n=15, k=3, p=0.3, seed=42)` and relabeled from `P1` through `P15`. The realization contains {ppi['nodes']} nodes and {ppi['edges']} edges, has density {ppi['density']:.4f}, average degree {ppi['average_degree']:.2f}, and diameter {ppi['diameter']}. Degrees range from {ppi['minimum_degree']} to {ppi['maximum_degree']}. Connectivity was verified as **{ppi['connected']}**.

#### 8.2.3 Layout procedure and criteria

Readability was assessed through edge crossings, label separation, preservation of local ring structure, visibility of rewired shortcuts, and correspondence between graph distance and visual proximity. Layout appearance was not accepted as evidence of community structure.

#### 8.2.4 Verified visual result

![Six-layout comparison](../visualizations/static/g_ppi_layout_comparison.png)

*Figure 4. Six layouts of the identical synthetic `G_ppi` graph.*

Kamada-Kawai was selected for this realization because pairwise graph-distance stress was minimized and ring-like locality remained legible. Spring placement provided a close alternative. Circular and shell layouts provided orderly labeling but imposed geometry unrelated to distance optimization. Spectral placement emphasized eigenvector structure, while random placement provided no topology-aware organization.

#### 8.2.5 Critical interpretation

The measured average clustering coefficient was **{ppi['average_clustering']:.3f}**. Triangle-based clustering was absent in this exact seeded realization despite use of a Watts-Strogatz generator. The figure therefore supports interpretation of locality and shortcuts, but it does not support a claim that visual communities or triangular clusters were discovered. This distinction demonstrates why layout inspection must be checked against calculated statistics.

#### 8.2.6 Limitations and verification evidence

Layout preference remains dependent on graph size, parameterization, task, and labeling requirements. Six individual figures and a combined comparison were saved under `visualizations/static/`. Graph order, size, connectivity, edge count, degree range, diameter, and clustering were recalculated from the generated object.

### 8.3 Exercise 3: Student-course bipartite network

#### 8.3.1 Objective and data design

A two-mode network was constructed so that enrollment could be represented without inventing student-to-student or course-to-course edges. The synthetic tables contain {bipartite['students']} students, {bipartite['courses']} courses, and {bipartite['enrollments']} enrollment records.

#### 8.3.2 Graph construction, analytical procedure, and validation

Students were assigned to partition 0 and courses to partition 1. An undirected edge was added only when an enrollment row linked the two partitions. Bipartiteness was verified through `nx.algorithms.bipartite.is_bipartite`, which returned **{bipartite['is_bipartite']}**. The realized bipartite density was {bipartite['bipartite_density']:.4f}.

#### 8.3.3 Degree results

**Table 3. Enrollment degree by course**

| Course | Student enrollments |
|---|---:|
{course_rows}

An average of {bipartite['mean_courses_per_student']:.2f} courses was assigned per student, while an average of {bipartite['mean_students_per_course']:.2f} students was assigned per course. **{bipartite['most_popular_course']}** had the largest course degree ({bipartite['most_popular_course_enrollments']}), and **{bipartite['most_enrolled_student']}** had the largest student degree ({bipartite['most_enrolled_student_courses']}).

#### 8.3.4 Visual result

![Student-course graph](../visualizations/static/student_course_bipartite_graph.png)

*Figure 5. Synthetic enrollment network with student circles and course squares.*

The two-column arrangement made partition membership explicit. Node shape and color distinguished entity type, and degree differences were visible through incident enrollment edges.

#### 8.3.5 Academic interpretation

{bipartite['pattern']} Degree answered a different question in each partition: student degree represented course load, whereas course degree represented synthetic popularity. A one-mode projection was not used because derived ties were not required.

#### 8.3.6 Limitations and verification evidence

The names and enrollment records were generated solely for teaching. No conclusion was drawn about actual students or enrollment behavior. Source tables were saved under `data/generated/`, calculated degrees were saved in `outputs/tables/student_course_degree_summary.csv`, and partition validity is checked by an automated test.

### 8.4 Exercise 4: Weighted Barabasi-Albert visualization

#### 8.4.1 Objective and graph construction

The effect of edge-width encoding was examined on the supplied 100-node Barabasi-Albert graph. The topology was generated with `n=100`, `m=2`, and seed 42. It contains {weighted['nodes']} nodes and {weighted['edges']} edges, with average degree {weighted['average_degree']:.2f} and maximum degree {weighted['maximum_degree']}.

#### 8.4.2 Weight assignment procedure

Every edge received one reproducible integer weight from 1 through 10 through a separately seeded pseudorandom generator. The topology was copied before weighting, so the unweighted and weighted figures contain identical nodes and edges. Edge width was scaled from weight; degree and layout position were not recomputed from weight.

#### 8.4.3 Verified weight distribution

**Table 4. Frequency of assigned edge weights**

| Edge weight | Number of edges |
|---:|---:|
{weight_rows}

The range was {weighted['minimum_weight']}-{weighted['maximum_weight']}, the mean was {weighted['mean_weight']:.3f}, the median was {weighted['median_weight']:.1f}, and the population standard deviation was {weighted['weight_standard_deviation']:.3f}. Total assigned edge weight was {weighted['total_edge_weight']:,}.

#### 8.4.4 Visual result

![Weighted BA graph](../visualizations/static/barabasi_albert_weighted.png)

*Figure 6. Seeded edge weights represented by line width in the synthetic Barabasi-Albert graph.*

Thicker lines increased perceptual salience for high-weight edges, while the hub-and-spoke topology remained unchanged. The unweighted companion image permits a direct encoding comparison.

#### 8.4.5 Academic interpretation

Edge weight and node centrality were kept conceptually separate. A high-weight edge was not interpreted as a high-degree node, and a prominent hub was not assumed to possess high-weight incident ties. Weighted centrality would require the weight to be defined as strength, capacity, cost, or distance; such a definition was not justified for the random teaching weights.

#### 8.4.6 Limitations and verification evidence

The weights are synthetic and demonstrate encoding only. The complete edge table was saved in `outputs/tables/barabasi_albert_edge_weights.csv`. Automated checks verify node and edge counts, the 1-10 range, and deterministic assignment.

### 8.5 Exercise 5: Statistical comparison of generative models

#### 8.5.1 Objective, model construction, and parameters

Three mechanisms were compared under the supplied parameters: Erdos-Renyi G(30, 0.08), Watts-Strogatz WS(30, 4, 0.1), and Barabasi-Albert BA(30, 2), each with seed 42. Random mixing, local clustering with rewiring, and preferential attachment were thereby contrasted.

#### 8.5.2 Metric procedure

Node count, edge count, density, average degree, clustering, connectivity, component count, average shortest-path length, maximum degree, and degree standard deviation were calculated. For the disconnected Erdos-Renyi realization, path length was calculated only on the largest connected component; a misleading whole-graph value was not reported.

#### 8.5.3 Verified statistical results

**Table 5. Statistical comparison of the three seeded models**

| Model | Nodes | Edges | Density | Avg. degree | Clustering | Avg. path | Components | Max degree | Degree SD |
|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|
{model_rows}

![Generative-model degree distributions](../visualizations/static/generative_models_degree_distribution.png)

*Figure 7. Degree distributions for the three exact seeded realizations.*

#### 8.5.4 Model-by-model interpretation

The Erdos-Renyi realization produced three components, low clustering ({models['records'][0]['Average Clustering']:.3f}), and maximum degree {models['records'][0]['Maximum Degree']}; it served as a homogeneous random baseline. The Watts-Strogatz realization remained connected and produced the highest clustering ({models['records'][1]['Average Clustering']:.3f}) with a narrow degree distribution. The Barabasi-Albert realization remained connected, produced the shortest average path ({models['records'][2]['Average Shortest Path']:.3f}), and generated maximum degree {models['records'][2]['Maximum Degree']} with degree standard deviation {models['records'][2]['Degree Standard Deviation']:.3f}.

#### 8.5.5 Comparative conclusion

{models['conclusion']} The small-world mechanism was better suited to local clustering and short-path interpretation, whereas preferential attachment was better suited to hub and degree-heterogeneity interpretation. Social-network realism was therefore treated as multidimensional.

#### 8.5.6 Limitations and verification evidence

Only one small seeded realization was compared per parameter set. Ranking may change across seeds, sizes, and parameters; repeated simulation with uncertainty intervals would be required for population-level inference. The comparison was saved in `outputs/tables/generative_models_comparison.csv`, and separate degree-distribution figures were generated for every model.

### 8.6 Exercise 6: Interactive node-color dashboard

#### 8.6.1 Objective and interaction design

An interactive view was produced so that one directed social graph could be interpreted through two node-color variables without a positional confound. Categorical color represents interest group, while a continuous Viridis scale represents in-degree.

#### 8.6.2 Graph construction and generation procedure

The supplied seeded logic produced {social['nodes']} users and {social['edges']} directed follow edges. Directed density was {social['density']:.4f}. Interest-group similarity increased selection probability during generation, and duplicate selections were removed before edge insertion.

#### 8.6.3 Group and degree results

**Table 6. Synthetic interest-group membership**

| Interest group | Users |
|---|---:|
{group_rows}

Mean in-degree was {social['mean_in_degree']:.2f}. In-degree ranged from {social['minimum_in_degree']} to {social['maximum_in_degree']}, and out-degree ranged from {social['minimum_out_degree']} to {social['maximum_out_degree']}. The three largest in-degree results were {top_users}.

#### 8.6.4 Dashboard implementation and verification

The dropdown contains **Interest Group** and **In-Degree** modes. Node coordinates, edges, labels, and hover fields are preserved; only marker color, legend, and scale metadata are changed. Hover text reports identifier, group, in-degree, and out-degree. The stored position-preservation result is **{social['positions_preserved']}**.

![Network dashboard colored by interest group](../screenshots/network-toggle-before.png)

*Figure 8. Initial dashboard state with categorical interest-group color.*

![Network dashboard colored by in-degree](../screenshots/network-toggle-in-degree.png)

*Figure 9. Continuous in-degree color with unchanged node positions.*

#### 8.6.5 Academic interpretation

The categorical mode supports comparison of group mixing, while the continuous mode emphasizes incoming-tie popularity. Because position is fixed, perceptual change can be attributed to color rather than a new layout. The exercise demonstrates coordinated encoding rather than evidence about real social-media users.

#### 8.6.6 Limitations and verification evidence

The graph is synthetic. Arrowheads are not displayed in the Plotly line traces, so direction is conveyed through calculated degrees rather than edge-end markers. The artifact was saved as `{social['standalone_html']}`, and both dropdown states were verified through browser testing.

### 8.7 Exercise 7: Research-domain knowledge graph

#### 8.7.1 Objective and semantic schema

A synthetic knowledge graph was constructed for Applied AI and Multimedia. Nodes were typed as research areas, methods, tools, applications, or outcomes. Edges received a named semantic relationship and an illustrative strength from 1 through 3.

#### 8.7.2 Graph construction and metric procedure

The graph contains {domain['nodes']} nodes and {domain['edges']} undirected links. Density is {domain['density']:.4f}, average degree is {domain['average_degree']:.2f}, and connectivity was verified as **{domain['connected']}**. Degree centrality, inverse-strength weighted betweenness, closeness, and strength-weighted PageRank were calculated. Edge distance was defined as (1 / strength) before shortest-path betweenness was evaluated.

#### 8.7.3 Verified centrality results

**Table 7. Five highest weighted-betweenness nodes**

| Node | Type | Degree | Betweenness | Closeness | PageRank |
|---|---|---:|---:|---:|---:|
{centrality_rows}

**{domain['top_betweenness_node']}** received the largest weighted-betweenness value ({domain['top_betweenness_value']:.4f}).

#### 8.7.4 Visual result

![Research-domain graph](../visualizations/static/domain_graph.png)

*Figure 10. Synthetic Applied AI and Multimedia knowledge graph with typed nodes and weighted links.*

#### 8.7.5 Academic interpretation

{domain['interpretation']} Centrality remained dependent on the modeled relationships: a method or application may appear prominent because otherwise separated categories are connected through it.

#### 8.7.6 Limitations and verification evidence

The nodes, relationships, and strengths were designed for demonstration and were not extracted from publications. The inverse-strength transformation assumes that a stronger semantic relationship represents a shorter effective distance; a different substantive meaning for strength would require a different transformation. Tables were saved under `data/generated/` and `outputs/tables/domain_graph_metrics.csv`, and both static and interactive visualizations were generated.
"""


def _extended_academic_sections(
    summary: dict[str, Any],
) -> list[tuple[str, list[str], None, None]]:
    """Return shared critical context for Markdown, DOCX, and PDF reports."""

    engagement = summary["engagement"]
    audit = summary["data_cleaning"]
    exercises = summary["exercises"]
    return [
        (
            "7.1 Conceptual framework for engagement measurement",
            [
                (
                    "Engagement is treated as a family of observable platform actions "
                    "rather than a direct measurement of attention, satisfaction, or "
                    "persuasion. A reaction can be produced with comparatively little "
                    "effort, whereas a comment or share may require additional cognitive "
                    "or social commitment. Even this distinction is contextual: a short "
                    "comment may be less consequential than a reaction from an influential "
                    "account, and a share may be critical, supportive, or merely archival. "
                    "The available table contains counts but no semantic evidence about "
                    "motivation. Accordingly, the analysis preserves reactions, comments, "
                    "and shares as separate variables before combining them into total "
                    "engagement. The aggregate is useful for ranking overall activity, but "
                    "it is not presented as a validated psychological scale."
                ),
                (
                    "The difference between exposure and response is also fundamental. "
                    "An engagement rate normally requires a denominator such as reach, "
                    "impressions, followers, or unique viewers. None is available in the "
                    "source file. Dividing observed actions by an invented or inappropriate "
                    "denominator would create false precision, so the project reports "
                    "counts, composition ratios, and within-dataset comparisons instead. "
                    "This constraint narrows the claims but improves their validity. It "
                    "also explains why a post with many engagements cannot automatically "
                    "be described as efficient: the number of people who had an opportunity "
                    "to engage is unknown."
                ),
            ],
            None,
            None,
        ),
        (
            "7.2 Evidence standards and reproducible design",
            [
                (
                    "The workflow follows a separation between source evidence, derived "
                    "evidence, and presentation. The raw CSV is preserved and identified "
                    "by SHA-256. Cleaning rules operate on a copy and produce an auditable "
                    "processed table. Calculated statistics and graph metrics are then "
                    "serialized in one machine-readable summary. Reports, notebook claims, "
                    "and website text read from that summary instead of maintaining "
                    "independent handwritten numbers. This design reduces the risk that "
                    "a corrected metric remains stale in one publication surface. It also "
                    "makes a discrepancy testable: a reported value can be traced to a "
                    "specific JSON field and to the function that calculated it."
                ),
                (
                    "Reproducibility is understood as more than the presence of code. "
                    "Paths are resolved from the repository root, stochastic operations "
                    "use seed 42, graph parameters are recorded, notebooks are executed "
                    "from clean kernels, and required artifacts are checked for existence "
                    "and nontrivial size. Each mandatory exercise is published as a "
                    "Black-formatted Python source file, a clean notebook, a clean-kernel "
                    "executed notebook, and a standalone browser notebook; interactive "
                    "exercises also retain their Plotly HTML. The website exposes these "
                    "files beside the exercise evidence, while the GitHub repository keeps "
                    "their canonical organized copies. The artifact manifest records hashes "
                    "for the canonical data, tables, figures, interactive files, notebooks, "
                    "Python sources, and reports. These measures cannot guarantee that every "
                    "software version will render pixels identically, but they make the "
                    "analytical decisions and the delivered version observable and comparable."
                ),
                (
                    "The design also separates regeneration from verification. A pipeline "
                    "can finish without proving that its outputs are correct, so completion "
                    "is followed by independent checks: schema assertions, numerical "
                    "cross-checks, notebook error inspection, document-structure audits, "
                    "link validation, responsive browser review, and deployment status "
                    "inspection. Failures are reported at the surface where they occur. "
                    "For example, an analytical test cannot certify that a dropdown works, "
                    "and a visual browser check cannot establish that a matrix used the "
                    "correct edge-weight field. Layered verification prevents one successful "
                    "tool from being treated as evidence for an unrelated requirement."
                ),
            ],
            None,
            None,
        ),
        (
            "7.3 Data quality assessment and transformation rationale",
            [
                (
                    f"The raw table contains {audit['initial_rows']:,} records and "
                    f"{audit['initial_columns']} columns. Four columns are empty across "
                    "all rows and carry no analytical information; removing them is a "
                    "structural correction rather than an imputation decision. Duplicate "
                    "checks are performed both on complete rows and on `status_id`, because "
                    "two identical records and two records sharing an identifier represent "
                    "different quality concerns. The observed duplicate counts are "
                    f"{audit['duplicate_rows_removed']} complete rows and "
                    f"{audit['duplicate_ids_removed']} identifiers. Date parsing is "
                    "validated before temporal features are produced, preventing malformed "
                    "timestamps from silently entering hourly or monthly summaries."
                ),
                (
                    "Count variables are required to be numeric, finite, and nonnegative. "
                    "Categorical labels are normalized so that spelling or capitalization "
                    "does not split a post type into artificial groups. Missing-value "
                    "handling is deliberately conservative: imputation is applied only "
                    "where a defensible zero interpretation exists, and the audit records "
                    "the number of affected values. The processed table is then validated "
                    "against explicit invariants, including unique identifiers, complete "
                    "engineered columns, nonnegative engagement, and finite ratios. These "
                    "checks convert assumptions into executable conditions rather than "
                    "leaving them as undocumented analyst judgment."
                ),
                (
                    f"Extreme engagement is retained. The upper Tukey fence is "
                    f"{audit['outlier_threshold']:,.2f}, and {audit['outlier_count']:,} "
                    "records exceed it. Deleting these observations would remove genuine "
                    "high-performing posts and materially change the phenomenon under "
                    "study. Instead, an outlier flag supports sensitivity-aware summaries, "
                    "while log transformation makes distribution plots readable. Mean "
                    "and median are reported together so that the influence of the long "
                    "right tail remains visible."
                ),
            ],
            None,
            None,
        ),
        (
            "7.4 Feature validity and analytical interpretation",
            [
                (
                    "Total engagement is defined as the sum of the dataset's authoritative "
                    "reaction total, comments, and shares. Component reaction fields are "
                    "retained for composition analysis and quality comparison, but they "
                    "do not replace the documented total without evidence that the two "
                    "constructs are identical. The comment-to-reaction and share-to-total "
                    "ratios are calculated with a zero-safe operation. Returning zero when "
                    "the denominator is zero avoids infinite values and preserves a clear "
                    "interpretation: no observed denominator-supported ratio is available "
                    "for that row."
                ),
                (
                    "Temporal features translate timestamps into posting hour, weekday, "
                    "month, ISO week, weekend status, and broad time-of-day categories. "
                    "They describe when posts were published, not when engagement occurred. "
                    "A high median for a posting hour may reflect content selection, seller "
                    "strategy, seasonal campaigns, or audience composition. The feature "
                    "therefore supports descriptive scheduling analysis but not a causal "
                    "recommendation that changing the clock time alone will improve results."
                ),
                (
                    "The low, medium, and high engagement bands are rank-based descriptive "
                    "labels within this dataset. They are not universal performance "
                    "benchmarks and should not be compared directly with another page or "
                    "platform without recalibration. Similarly, the outlier indicator is "
                    "a statistical flag rather than a judgment that a record is erroneous. "
                    "These labels are useful for filtering and visual explanation only "
                    "when their local definition remains visible."
                ),
            ],
            None,
            None,
        ),
        (
            "7.5 Distribution-aware exploratory strategy",
            [
                (
                    f"The {engagement['rows']:,} posts record "
                    f"{engagement['total_engagement']:,} total engagements. The mean "
                    f"({engagement['mean_engagement']:,.2f}) is substantially larger than "
                    f"the median ({engagement['median_engagement']:,.2f}), demonstrating "
                    "that a small proportion of posts contributes heavily to the aggregate. "
                    "A mean-only comparison would therefore describe the economic weight "
                    "of high observations but not the typical post. The report pairs means "
                    "with medians, count distributions, and log-scaled views so that both "
                    "questions remain answerable."
                ),
                (
                    "Post-type comparisons are observational. Video has the highest "
                    f"observed mean total engagement ({engagement['top_post_type_mean_engagement']:,.2f}), "
                    "but post type was not randomized. Seller identity, live-event context, "
                    "promotion, topic, audience size, and the platform's historical ranking "
                    "system may all affect the result. The finding is appropriately phrased "
                    "as an association in the recorded sample. It can motivate a controlled "
                    "content experiment, but it cannot substitute for one."
                ),
                (
                    "Spearman correlation is selected because it evaluates monotonic rank "
                    "association without assuming normally distributed raw counts. Total "
                    "engagement necessarily shares components with reactions, comments, "
                    "and shares, so strong correlations with the total are partly "
                    "mathematical. Correlation between components may also arise from a "
                    "common exposure mechanism. The heatmap is therefore interpreted as "
                    "a descriptive relationship map, not as evidence that one action causes "
                    "another."
                ),
            ],
            None,
            None,
        ),
        (
            "7.6 Graph representation and metric semantics",
            [
                (
                    "A graph is defined by its nodes, edges, direction, and attributes; "
                    "the drawing is only one representation of that structure. Adjacency "
                    "lists are efficient for neighbor inspection in sparse graphs, while "
                    "matrices support exact pairwise lookup, linear-algebra operations, "
                    "and symmetry tests. Node-link diagrams are effective for paths, hubs, "
                    "bridges, and small-scale topology, but they become difficult to read "
                    "as density and label count increase. The exercises use each form for "
                    "the task it supports rather than treating visualization as decoration."
                ),
                (
                    "Metric meaning depends on the graph model. Degree counts incident "
                    "ties, in-degree counts received directed ties, and PageRank distributes "
                    "importance through linked neighbors. Betweenness identifies nodes on "
                    "shortest paths, but weighted shortest paths require a distance. In the "
                    "knowledge graph, semantic strength is transformed to inverse strength "
                    "so that a strong relationship becomes a short effective distance. "
                    "Using raw strength directly as distance would reverse the intended "
                    "meaning and could change the ranking."
                ),
                (
                    "Connectivity also controls path interpretation. Average shortest path "
                    "is undefined across disconnected components because no finite route "
                    "joins every pair. The model comparison therefore reports whole-graph "
                    "path length only for connected graphs and uses the largest connected "
                    "component for the disconnected Erdos-Renyi realization. The scope is "
                    "stored beside the value so that a reader cannot mistake the component "
                    "result for a network-wide statistic."
                ),
            ],
            None,
            None,
        ),
        (
            "7.7 Visual encoding and layout validity",
            [
                (
                    "Layout algorithms change coordinates, not topology. A spring layout "
                    "uses attractive and repulsive forces, Kamada-Kawai minimizes a "
                    "graph-distance stress objective, spectral layout uses eigenvectors, "
                    "and circular, shell, or random placement applies different geometric "
                    "constraints. Because the same graph can look clustered or dispersed "
                    "under different coordinates, visual grouping must be checked against "
                    "calculated community or clustering measures. In Exercise 2, node "
                    "attributes, sizes, edge styles, labels, and figure dimensions are "
                    "held constant so that placement is the controlled variable."
                ),
                (
                    f"The supplied `G_ppi` realization has average clustering "
                    f"{exercises['exercise_2']['average_clustering']:.3f}. Kamada-Kawai "
                    "is selected because it gives the clearest distance-oriented account "
                    "of ring locality and shortcuts, not because it reveals communities "
                    "that the metric does not support. This distinction is important in "
                    "academic graph drawing: an aesthetically separated region may be a "
                    "layout artifact, while a visually dense region may reflect label "
                    "placement rather than structural cohesion."
                ),
                (
                    "Color, size, shape, and width are assigned specific semantic roles. "
                    "Categorical variables use distinct hues; ordered measures use a "
                    "continuous scale; student and course partitions use separate node "
                    "forms; and synthetic edge weight is encoded through width. Legends, "
                    "captions, and hover text state these mappings. Redundant visual cues "
                    "are used where they improve accessibility, while unnecessary three-"
                    "dimensional effects and decorative encodings are avoided."
                ),
            ],
            None,
            None,
        ),
        (
            "7.8 Static and interactive visualization roles",
            [
                (
                    "Static figures provide a stable scholarly record. They can be "
                    "numbered, cited, printed, and compared without requiring a runtime. "
                    "They are exported at high resolution with explicit titles, labels, "
                    "legends, and captions. Interactive Plotly artifacts answer different "
                    "needs: readers can inspect individual nodes and posts, zoom into the "
                    "long tail, isolate categories, and switch encodings. The project "
                    "therefore treats interactivity as an analytical affordance rather "
                    "than a replacement for documented static evidence."
                ),
                (
                    "The node-color dashboard illustrates controlled interaction. Both "
                    "modes retain the same coordinates and edge traces. Interest-group "
                    "color is categorical, whereas in-degree color is continuous. Because "
                    "geometry remains fixed, changes in perception can be attributed to "
                    "the selected color variable rather than a simultaneous layout change. "
                    "Hover fields disclose group, in-degree, and out-degree, enabling the "
                    "reader to check the visual impression against exact values."
                ),
            ],
            None,
            None,
        ),
        (
            "12.1 Critical synthesis of empirical and synthetic evidence",
            [
                (
                    "The project contains two related but noninterchangeable forms of "
                    "evidence. The Facebook analysis is empirical at the post level: rows "
                    "are observed records and calculated values describe that sample. The "
                    "network exercises are synthetic demonstrations of graph construction, "
                    "measurement, and visual encoding. Their reproducibility does not make "
                    "them empirical claims about Facebook users, students, transport, "
                    "proteins, or research communities. Maintaining this boundary protects "
                    "the report from a visually persuasive but unsupported narrative."
                ),
                (
                    "Taken together, the components show why data visualization requires "
                    "both numerical and perceptual validation. A distribution plot explains "
                    "why the mean and median differ; a matrix test verifies a symmetry that "
                    "a node-link drawing merely suggests; clustering corrects an overly "
                    "enthusiastic visual reading of layout; and a fixed-coordinate dropdown "
                    "isolates the effect of color. Each major interpretation is therefore "
                    "paired with a calculated value, controlled encoding, or explicit "
                    "limitation."
                ),
            ],
            None,
            None,
        ),
        (
            "12.2 Generalizability and uncertainty",
            [
                (
                    "External validity is limited by time, geography, industry, and platform "
                    "context. The posts end in 2018 and originate from ten Thai fashion "
                    "and cosmetics sellers. Facebook interfaces, ranking systems, audience "
                    "behavior, and commercial practices have changed since collection. "
                    "Results should therefore be read as a careful description of this "
                    "dataset, not as a current global benchmark for social-media strategy."
                ),
                (
                    "The graph-model comparison is also conditional. Metrics from one "
                    "30-node seeded realization are exact for that realization but uncertain "
                    "as estimates of a model's typical behavior. A stronger simulation "
                    "study would repeat each parameter set across many seeds, summarize "
                    "the metric distributions, and report confidence or percentile "
                    "intervals. Parameter sweeps could then separate a mechanism's general "
                    "tendency from a particular random draw."
                ),
                (
                    "Sensitivity analysis would also strengthen the empirical component. "
                    "Post-type rankings could be recalculated after winsorization, within "
                    "time periods, or with seller-level controls if stable page identifiers "
                    "became available. Median and quantile regression could describe "
                    "different parts of the engagement distribution without allowing the "
                    "largest posts to dominate. Missing exposure denominators remain the "
                    "most important obstacle: no transformation of the current counts can "
                    "recover the unobserved population that saw each post. Future inference "
                    "should therefore improve data collection before increasing model "
                    "complexity."
                ),
            ],
            None,
            None,
        ),
        (
            "12.3 Ethical, privacy, and academic-integrity safeguards",
            [
                (
                    "The source is anonymized and contains no post text, account names, or "
                    "direct personal identifiers. The project does not attempt linkage, "
                    "re-identification, or inference about protected characteristics. "
                    "Aggregated reporting further reduces exposure. Synthetic people and "
                    "relationships are labeled as such in data files, figures, notebooks, "
                    "the website, and the report. These labels are not cosmetic: they "
                    "prevent readers from confusing pedagogical structures with observed "
                    "human behavior."
                ),
                (
                    "Academic integrity is supported through source attribution, explicit "
                    "acquisition history, calculated rather than fabricated statistics, "
                    "and reproducible code. The UCI dataset and its accompanying article "
                    "are cited, software tools are acknowledged, and the unlicensed "
                    "teaching package is not republished. The report distinguishes a "
                    "Kaggle mirror from the actual UCI acquisition route and does not claim "
                    "that credentials or an API were used when they were unavailable."
                ),
                (
                    "Accessibility is treated as an ethical property of communication. "
                    "Figures use descriptive captions and alternative text; tables identify "
                    "header rows; keyboard focus remains visible; and color is accompanied "
                    "by labels, shape, position, or downloadable values when feasible. "
                    "Interactive frames retain meaningful titles, while static equivalents "
                    "preserve the central result for readers who cannot or prefer not to "
                    "operate a dynamic chart. These measures do not constitute a complete "
                    "accessibility certification, but they reduce avoidable barriers and "
                    "make the evidential content less dependent on a single sensory channel."
                ),
            ],
            None,
            None,
        ),
        (
            "12.4 Practical recommendations",
            [
                (
                    "For applied content analysis, future data collection should prioritize "
                    "reach and impression denominators, stable page identifiers, content "
                    "metadata, and campaign context. These variables would permit rate "
                    "estimation, multilevel seller comparisons, and adjustment for exposure. "
                    "A prospective design could rotate content formats or posting windows "
                    "under controlled conditions. Until such evidence exists, video should "
                    "be described as the highest-engagement observed category rather than "
                    "a guaranteed intervention."
                ),
                (
                    "For network analysis, repeated simulations should precede general "
                    "claims about model behavior. Community detection should be calculated "
                    "independently of layout, and weighted metrics should document whether "
                    "an edge attribute represents strength, capacity, probability, cost, "
                    "or distance. Interactive views should preserve a stable reference "
                    "state when comparing visual variables and should expose exact values "
                    "through accessible text or downloadable tables."
                ),
                (
                    "For project maintenance, the canonical summary and artifact manifest "
                    "should remain the only numerical publication interfaces. Changes to "
                    "cleaning, graph parameters, or metric definitions should trigger the "
                    "complete pipeline, notebook execution, report generation, website "
                    "build, and cross-surface tests. A published release should be tied "
                    "to a commit SHA so that the local project, repository, and GitHub "
                    "Pages artifact can be compared unambiguously."
                ),
                (
                    "For teaching and assessment, each exercise should remain inspectable "
                    "at three levels. The notebook shows how the calculation unfolds, the "
                    "standalone source module shows reusable implementation, and the website "
                    "connects code to output and interpretation. This layered presentation "
                    "supports readers with different technical backgrounds while preserving "
                    "one analytical result. Assessment should reward the justification of "
                    "a representation and the accuracy of its limitation, not only the "
                    "presence of an attractive graph."
                ),
            ],
            None,
            None,
        ),
        (
            "12.5 Integrated design framework",
            [
                (
                    "The project adopts a question-first visualization framework. The "
                    "empirical questions concern differences in recorded engagement across "
                    "post types and time; the exercise questions concern graph structure, "
                    "layout, weighting, and interaction. Those question classes require "
                    "different evidence. Comparisons use aligned bars and tables, skewed "
                    "outcomes use distributions and robust summaries, pairwise monotonic "
                    "association uses a rank-correlation matrix, and network questions use "
                    "matrices, node-link views, and calculated graph metrics. Selecting a "
                    "representation from the analytical task prevents decorative novelty "
                    "from becoming the governing design criterion."
                ),
                (
                    "Network science also distinguishes structure from its visual "
                    "realization. Barabasi (2016) treats degree distributions, clustering, "
                    "paths, and generative mechanisms as properties of graphs rather than "
                    "of drawings. This report follows that distinction by calculating "
                    "metrics with NetworkX and using layouts only to expose selected "
                    "relationships. Hagberg, Schult, and Swart (2008) describe NetworkX as "
                    "a platform for structural analysis; its role here is therefore "
                    "computational, while Matplotlib and Plotly provide complementary "
                    "static and interactive presentations. Recalculating metrics before "
                    "rendering keeps visual appearance subordinate to the stored graph."
                ),
                (
                    "The empirical component is similarly anchored in provenance. Dehouche "
                    "(2018, 2020) documents the Facebook Live Sellers in Thailand data and "
                    "its engagement variables. The present study retains that observational "
                    "unit and does not reinterpret posts as users or infer social ties. It "
                    "adds reproducible cleaning, feature definitions, descriptive analysis, "
                    "and visual communication, but it does not enlarge the evidential scope "
                    "of the source. This source-to-claim traceability is especially important "
                    "when a project combines a real table with synthetic graph exercises: "
                    "the common software environment must not be mistaken for a common "
                    "population or a shared causal design."
                ),
                (
                    "Finally, the communication design recognizes that reproducibility is "
                    "both technical and rhetorical. A fixed seed and checksum reproduce an "
                    "artifact, but a reader also needs definitions, captions, limitations, "
                    "and accessible alternatives to understand what that artifact supports. "
                    "The notebooks expose calculations, the report supplies sustained "
                    "argument, the website supports navigation and interaction, and the "
                    "machine-readable summary binds their numerical claims. Agreement among "
                    "these layers is tested rather than assumed. The resulting architecture "
                    "makes revision auditable: a change to a statistic must propagate from "
                    "the canonical JSON through every presentation surface or cause a "
                    "contract test to fail."
                ),
            ],
            None,
            None,
        ),
    ]


def _sections_as_markdown(
    sections: list[tuple[str, list[str], None, None]],
) -> str:
    return "\n\n".join(
        f"### {title}\n\n" + "\n\n".join(paragraphs)
        for title, paragraphs, _, _ in sections
    )


def build_markdown(summary: dict[str, Any]) -> str:
    """Create a complete academic report in Markdown from calculated results."""

    engagement = summary["engagement"]
    audit = summary["data_cleaning"]
    transport = summary["exercises"]["exercise_1"]
    bipartite = summary["exercises"]["exercise_3"]
    weighted = summary["exercises"]["exercise_4"]
    models = summary["exercises"]["exercise_5"]
    domain = summary["exercises"]["exercise_7"]
    type_rows = "\n".join(
        (
            f"| {row['status_type'].title()} | {row['posts']:,} | "
            f"{row['mean_reactions']:,.2f} | {row['mean_comments']:,.2f} | "
            f"{row['mean_shares']:,.2f} | {row['mean_total_engagement']:,.2f} |"
        )
        for row in engagement["by_type"]
    )
    model_rows = "\n".join(
        (
            f"| {row['Model']} | {row['Nodes']} | {row['Edges']} | "
            f"{row['Density']:.3f} | {row['Average Clustering']:.3f} | "
            f"{row['Average Shortest Path']:.3f} ({row['Path Length Scope']}) | "
            f"{row['Maximum Degree']} | {row['Degree Standard Deviation']:.3f} |"
        )
        for row in models["records"]
    )
    report = f"""# Visual Analytics and Network Analysis of Facebook Engagement Using Python

**Student:** {STUDENT['name']}

**Student ID:** {STUDENT['id']}

**Course:** {STUDENT['course_name']} ({STUDENT['course_code']})

**Semester:** {STUDENT['semester']}

**Teacher:** {STUDENT['teacher']}, {STUDENT['designation']}

**Department:** {STUDENT['department']}

**University:** {STUDENT['university']}

## Abstract

An analysis is presented of {engagement['rows']:,} anonymized Facebook posts from ten Thai fashion and cosmetics sellers, and the supplied graph-visualization teaching material is extended through seven reproducible exercises. Data cleaning, feature engineering, exploratory visualization, graph representations, layout comparison, bipartite modeling, weighted encoding, generative-model statistics, and interactive Plotly networks are combined. A total of {engagement['total_reactions']:,} reactions, {engagement['total_comments']:,} comments, and {engagement['total_shares']:,} shares is recorded. The largest mean total engagement is observed for {engagement['top_post_type_by_mean_engagement']} posts ({engagement['top_post_type_mean_engagement']:,.2f}), although a causal claim is not supported by the observational design. Symmetry is verified in the weighted `G_transport` adjacency matrix, while a trade-off is revealed by the model comparison: clustering and short paths are captured by Watts-Strogatz, whereas hubs and degree heterogeneity are captured by Barabasi-Albert. A reproducible academic artifact is produced with static figures, interactive HTML, executable notebooks, a responsive website, automated tests, and machine-readable analytical outputs.

**Keywords:** Facebook engagement, data visualization, social-network analysis, NetworkX, Plotly, graph layout, Python

## Contents

1. Introduction and research questions
2. Dataset and provenance
3. Tools and reproducibility
4. Data-cleaning methodology
5. Feature engineering
6. Exploratory data analysis
7. Graph representation, visualization, and evidence framework
8. Seven mandatory network exercises
9. Discussion
10. Limitations and ethical considerations
11. Conclusion and recommendations
12. Future work and critical implications
References and reproducibility appendices

## 1. Introduction

Social-media engagement is multidimensional: lightweight reactions coexist with comments and shares that may require greater effort. A responsible visual analysis must therefore preserve individual measures while also making aggregate patterns legible. This study addresses two linked tasks. First, it describes post-level engagement in an anonymized, publicly licensed Facebook dataset. Second, it applies and extends the graph concepts in the supplied course notebook so that representations, layouts, weights, centrality, and interaction can be compared using reproducible evidence.

### 1.1 Problem statement

Raw engagement counts are highly skewed, post categories are imbalanced, and a tabular engagement dataset does not contain verified user-to-user relationships. Treating rows as a social network would invent edges. The project therefore keeps the empirical Facebook analysis tabular and labels the assignment's relationship graphs as synthetic.

### 1.2 Objectives

1. Construct a reproducible cleaning and feature-engineering workflow.
2. Describe engagement differences across post type and time without causal overreach.
3. Complete all seven mandatory graph exercises using the supplied graph definitions.
4. Compare static and interactive visual encodings.
5. publish auditable outputs through documentation, tests, a report, and a responsive dashboard.

### 1.3 Research questions

1. How do reactions, comments, shares, and total engagement vary across post types?
2. Which temporal patterns are visible in posting activity and median engagement?
3. How do graph layout and edge-weight encodings change structural interpretation?
4. Which canonical generative model reproduces particular social-network properties?
5. Which nodes occupy structurally important positions in the synthetic research graph?

## 2. Dataset and provenance

The dataset is **Facebook Live Sellers in Thailand**, authored by Nassim Dehouche and distributed by the UCI Machine Learning Repository under CC BY 4.0 (Dehouche, 2018). The raw CSV contains 7,050 rows and 16 columns; four columns are entirely empty placeholders. Posts span {engagement['date_min'][:10]} through {engagement['date_max'][:10]}. The official UCI file was used because Kaggle credentials were not configured; a Kaggle mirror is documented but was not used for acquisition. The accompanying study reports anonymized data collected from ten Thai fashion and cosmetics seller pages (Dehouche, 2020).

The selection is appropriate because it is medium-sized, includes numerical and categorical engagement variables, contains timestamps, is publicly licensed, and contains no names or message text. Its limitation is equally important: it does not contain reach, impressions, follower counts, page identity, or user-to-user interaction IDs.

## 3. Tools and reproducibility

Python is the primary language. Pandas performs data preparation, NumPy supports numerical operations, NetworkX constructs and measures graphs, Matplotlib creates 300-DPI static figures, Plotly produces standalone interactive HTML, SciPy supports NetworkX's Kamada-Kawai layout path, python-docx creates the editable report, ReportLab provides a PDF fallback, nbformat creates notebooks, and pytest validates the workflow. Fixed seed 42 controls every stochastic graph and layout where the API accepts a seed.

## 4. Data-cleaning methodology

The raw file was preserved unchanged. Column names were standardized; {len(audit['empty_columns_removed'])} entirely empty columns were removed; numeric engagement fields were coerced, checked for nonnegative values, and stored as integers; timestamps were parsed; categories were normalized; and duplicates were checked by full row and `status_id`. The pipeline removed {audit['duplicate_rows_removed']} duplicate rows and {audit['duplicate_ids_removed']} duplicate IDs. It removed {audit['invalid_dates_removed']} invalid timestamps and imputed {audit['numeric_missing_values_imputed']} missing numeric values. The final analytical table contains {audit['final_rows']:,} rows and {audit['final_columns']} columns.

Outliers are flagged, not deleted, using the upper Tukey fence on total engagement ({audit['outlier_threshold']:,.2f}); {audit['outlier_count']:,} posts exceed the threshold. This preserves genuine high-performing posts while making skew explicit.

## 5. Feature engineering

- **Total engagement** = reactions + comments + shares.
- **Like-to-comment ratio** = likes / comments, with zero returned when comments are zero.
- **Share-to-engagement ratio** = shares / total engagement, with zero-safe division.
- **Positive reaction percentage** = (likes + loves + wows) / sum of recorded reaction components × 100.
- **Negative reaction percentage** = (sads + angrys) / sum of recorded reaction components × 100.
- **Temporal features** include posting hour, weekday, month, ISO week, weekend flag, and time-of-day band.
- **Engagement category** assigns low, medium, and high groups by rank-based tertiles.

No engagement-rate feature is claimed because reach or impressions are absent.

## 6. Exploratory data analysis

The dataset is composed of {engagement['post_type_counts']['photo']:,} photos, {engagement['post_type_counts']['video']:,} videos, {engagement['post_type_counts']['status']:,} status posts, and {engagement['post_type_counts']['link']:,} links. Total engagement of {engagement['total_engagement']:,} is recorded; a mean of {engagement['mean_engagement']:,.2f} and a median of {engagement['median_engagement']:,.2f} are obtained, indicating right skew. The largest observed median engagement is associated with {engagement['best_median_posting_hour']:02d}:00 ({engagement['best_median_posting_hour_value']:,.2f}), but the association may be affected by content mix, seller behavior, seasonality, or other unobserved factors.

**Table 1. Engagement summary by post type**

| Post type | Posts | Mean reactions | Mean comments | Mean shares | Mean total engagement |
|---|---:|---:|---:|---:|---:|
{type_rows}

![Mean engagement components by post type](../visualizations/static/engagement_by_post_type.png)

*Figure 1. Mean reactions, comments, and shares by post type.*

![Engagement distributions](../visualizations/static/engagement_distribution.png)

*Figure 2. Log-transformed total-engagement distributions reveal substantial right skew.*

![Correlation heatmap](../visualizations/static/engagement_correlation.png)

*Figure 3. Spearman correlations describe monotonic association rather than causation.*

## 7. Graph representation and visualization methodology

An adjacency matrix offers exact pairwise lookup, while a node-link diagram emphasizes paths, clusters, bridges, and hubs. Force-directed layouts encode graph topology into geometric proximity; circular and random layouts mainly impose placement rules unrelated to community quality. Node size, color, and edge width are kept semantically consistent within each comparison to isolate the effect of layout.

## 8. Mandatory hands-on network exercises

### 8.1 Exercise 1: Weighted adjacency representations

**Objective.** Reuse Section 8.2's seven-city `G_transport` graph and verify its weighted matrix.

**Method and implementation.** `nx.to_pandas_adjacency(..., weight="weight")` generated a labeled matrix, and `numpy.allclose(A, A.T)` tested symmetry.

**Result.** The matrix is symmetric: **{transport['matrix_symmetric']}**. Dhaka has six incident routes and therefore occupies the main hub row and column. The matrix stores kilometer values, not binary indicators.

**Interpretation.** An undirected edge contributes the same weight to `(u,v)` and `(v,u)`. The matrix would not need to be symmetric for a directed graph or when directional weights differ.

**Limitation.** The transportation graph is a teaching example rather than a complete national route network.

### 8.2 Exercise 2: Network layout comparison

**Objective.** Draw `G_ppi` with spring, circular, shell, spectral, Kamada-Kawai, and random layouts while holding visual encoding constant.

**Result and interpretation.** The **Kamada-Kawai layout** is selected as the clearest view for this small graph. It minimizes graph-theoretic distance stress across all node pairs and makes the ring-like locality plus rewired shortcuts more legible than circular, shell, spectral, or random placement. Spring is a close alternative. Crucially, the supplied `k=3` realization has a measured average clustering coefficient of {summary['exercises']['exercise_2']['average_clustering']:.3f}. No layout can reveal triangle-based clustering that is absent; the selected layout therefore clarifies distance structure and shortcuts, not empirical communities. The conclusion is layout-specific rather than a claim that one algorithm is universally superior.

![Six-layout comparison](../visualizations/static/g_ppi_layout_comparison.png)

*Figure 4. Six layouts of the identical `G_ppi` graph.*

**Limitation.** Visual cluster separation is not itself a community-detection result.

### 8.3 Exercise 3: Student-course bipartite network

**Objective.** Model synthetic enrollment between {bipartite['students']} students and {bipartite['courses']} courses from stored CSV tables.

**Result.** NetworkX confirms bipartiteness. The graph has {bipartite['enrollments']} enrollments. **{bipartite['most_popular_course']}** is most popular ({bipartite['most_popular_course_enrollments']} enrollments), and **{bipartite['most_enrolled_student']}** takes the most courses ({bipartite['most_enrolled_student_courses']}).

![Student-course graph](../visualizations/static/student_course_bipartite_graph.png)

*Figure 5. Synthetic enrollment graph; circles are students and squares are courses.*

**Interpretation.** {bipartite['pattern']}

**Limitation.** The graph is synthetic and cannot support conclusions about actual DIU enrollment.

### 8.4 Exercise 4: Weighted Barabasi-Albert visualization

**Objective.** Extend Section 6's 100-node BA graph with seeded integer edge weights from 1 through 10.

**Result.** {weighted['edges']} edges received weights with observed range {weighted['minimum_weight']}–{weighted['maximum_weight']} and mean {weighted['mean_weight']:.2f}. Width makes high-weight ties salient, but hubs remain defined by node degree rather than edge width.

![Weighted BA graph](../visualizations/static/barabasi_albert_weighted.png)

*Figure 6. Synthetic BA graph with normalized edge-width encoding.*

**Interpretation.** The view changes which ties attract attention, but random exercise weights do not constitute empirical importance. Edge weight and node centrality can interact, yet they are not interchangeable.

**Limitation.** Random weights demonstrate encoding only.

### 8.5 Exercise 5: Statistical comparison of generative models

**Table 2. Statistical comparison of Section 3 models**

| Model | Nodes | Edges | Density | Avg. clustering | Avg. path | Max degree | Degree SD |
|---|---:|---:|---:|---:|---:|---:|---:|
{model_rows}

![Generative-model degree distributions](../visualizations/static/generative_models_degree_distribution.png)

*Figure 7. Degree distributions expose homogeneous, small-world, and hub-dominated structures.*

**Interpretation.** Watts-Strogatz best captures high clustering with short paths, whereas Barabasi-Albert best captures hubs and heterogeneous degree. A realistic social network may combine both properties; none of these simple models is universally best. Erdős-Rényi provides a useful random baseline but lacks both mechanisms.

**Limitation.** Results come from one seeded 30-node realization per model and are sensitive to parameter choice.

### 8.6 Exercise 6: Interactive node-color dashboard

The standalone Plotly dashboard preserves node positions while a dropdown toggles between categorical interest-group color and continuous in-degree color. Hover text reports node ID, group, in-degree, and out-degree. Edges remain visible in both modes. The synthetic graph reuses the supplied Section 9 generation logic and is embedded in the final website.

**Limitation.** Directed edges are represented as line segments without arrowheads in Plotly; direction remains available through in/out-degree metrics and the graph definition.

### 8.7 Exercise 7: Research-domain network

**Objective.** Represent a synthetic knowledge graph for Applied AI and Multimedia using 15 typed nodes and 21 weighted semantic links.

**Result.** **{domain['top_betweenness_node']}** has the highest weighted betweenness ({domain['top_betweenness_value']:.4f}).

![Research-domain graph](../visualizations/static/domain_graph.png)

*Figure 8. Synthetic applied-AI and multimedia research knowledge graph.*

**Interpretation.** {domain['interpretation']}

**Limitation.** Nodes and relationships are illustrative and not extracted from publications.

## 9. Discussion

The empirical and synthetic components answer different questions. The Facebook table supports descriptive comparison of observed posts; it does not reveal individual users, verified follower ties, or causal effects. The synthetic graphs instead demonstrate how topology, representation, layout, and visual encoding affect interpretation. Keeping these scopes separate prevents a common analytical error: inferring social relationships that the source data never recorded.

The engagement analysis shows a mixture of high-frequency photos and high-engagement videos. Mean comparisons are influenced by extreme posts, so medians, log distributions, and Tukey flags accompany them. Rank correlations are preferable to a sole reliance on Pearson correlation under strong skew, but they still describe association only.

## 10. Limitations and ethical considerations

The data end in 2018, cover ten sellers in Thailand, and may not generalize to other countries, sectors, platforms, or current Facebook behavior. Page identity and content text are absent, reducing both explanatory power and privacy risk. Platform algorithms, audience size, campaign spend, and post quality are unobserved confounders. The analysis avoids re-identification, preserves anonymization, reports only aggregated results, and distinguishes real observations from synthetic exercises.

## 11. Conclusion and recommendations

A reproducible bridge between tabular engagement analytics and network visualization is demonstrated. The highest observed mean engagement is associated with video posts, but the finding is descriptive. The supplied small-world teaching graph is presented most clearly by Kamada-Kawai, tie salience is distinguished from node centrality in the weighted BA exercise, and the multidimensional nature of social-network realism is revealed by the generative-model comparison.

Recommended next steps are to: (1) obtain ethically collected reach and impression denominators; (2) analyze multiple pages and time periods; (3) use repeated simulations and uncertainty intervals when comparing graph models; (4) test community detection separately from visual layout; and (5) retain the distinction between observed edges and synthetic teaching structures.

## 12. Future work

Future work could incorporate content embeddings, causal designs, hierarchical models for seller-level heterogeneity, temporal network models when genuine relational identifiers exist, and accessibility evaluation with classroom users.

## References

Barabasi, A.-L. (2016). *Network science*. Cambridge University Press. http://networksciencebook.com/

Dehouche, N. (2018). *Facebook Live Sellers in Thailand* [Data set]. UCI Machine Learning Repository. https://doi.org/10.24432/C5R60S

Dehouche, N. (2020). Dataset on usage and engagement patterns for Facebook Live sellers in Thailand. *Data in Brief, 30*, 105661. https://doi.org/10.1016/j.dib.2020.105661

Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network structure, dynamics, and function using NetworkX. *Proceedings of the 7th Python in Science Conference*, 11–15.

Matplotlib Development Team. (2026). *Matplotlib documentation*. https://matplotlib.org/stable/

NetworkX Developers. (2026). *NetworkX documentation*. https://networkx.org/documentation/stable/

pandas development team. (2026). *pandas documentation*. https://pandas.pydata.org/docs/

Plotly Technologies Inc. (2026). *Plotly Python documentation*. https://plotly.com/python/

## Appendix A. Reproducibility and output map

Run `python main.py` to rebuild processed data, analytical tables, figures, interactive HTML, notebooks, summary JSON, and reports. Run `pytest` for automated validation. Each numbered exercise folder contains a standalone Python source, a clean notebook, a clean-kernel executed notebook, and a browser-ready HTML notebook; Exercises 6 and 7 also expose their interactive Plotly files. The raw CSV remains in `data/raw/`, the cleaned table in `data/processed/`, required exercise tables in `outputs/tables/`, figures in `visualizations/`, and interactive artifacts in both `visualizations/interactive/` and the deployed website.

## Appendix B. Assumptions

The supplied teaching notebooks contain source-identical graph definitions. Exercise 2 uses Kamada-Kawai as the preferred layout based on this exact realization. The Kaggle API was not used. `num_reactions` is treated as the dataset's authoritative reaction total while component sums are retained as a quality-control comparison.
"""
    extended = _extended_academic_sections(summary)
    pre_exercise = [section for section in extended if section[0].startswith("7.")]
    post_exercise = [section for section in extended if section[0].startswith("12.")]
    before_exercises, remaining = report.split(
        "## 8. Mandatory hands-on network exercises", maxsplit=1
    )
    _, after_exercises = remaining.split("## 9. Discussion", maxsplit=1)
    after_exercises = after_exercises.replace(
        "## References",
        _sections_as_markdown(post_exercise) + "\n\n## References",
        1,
    )
    return (
        before_exercises
        + _sections_as_markdown(pre_exercise)
        + "\n\n"
        + "## 8. Mandatory hands-on network exercises\n\n"
        + _exercise_markdown(summary).strip()
        + "\n\n## 9. Discussion"
        + after_exercises
    )


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    properties.append(table_header)


def _set_table_geometry(table, widths: list[float]) -> None:
    """Apply fixed 9360-DXA geometry and explicit cell margins."""

    dxa_widths = [round(width * 1440) for width in widths]
    if sum(dxa_widths) != 9360:
        raise ValueError(f"Table widths must total 6.5 inches: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": "9360", "w:type": "dxa"}),
        ("w:tblInd", {"w:w": "120", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        element = properties.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            properties.append(element)
        for name, value in attributes.items():
            element.set(qn(name), value)
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in dxa_widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, width in enumerate(dxa_widths):
            cell_width = row.cells[index]._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _normalize_docx_container(path: Path) -> None:
    """Give DOCX ZIP members fixed timestamps for reproducible hashes."""

    temporary = path.with_suffix(".normalized.docx")
    with (
        ZipFile(path) as source,
        ZipFile(temporary, "w", compression=ZIP_DEFLATED) as destination,
    ):
        for original in sorted(source.infolist(), key=lambda item: item.filename):
            info = copy(original)
            info.date_time = (2026, 7, 27, 0, 0, 0)
            info.compress_type = ZIP_DEFLATED
            destination.writestr(info, source.read(original.filename))
    temporary.replace(path)


def _configure_document(document: Document) -> None:
    """Apply the standard_business_brief preset with an editorial cover."""

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size, before, after, color in (
        ("Title", 30, 0, 8, NAVY),
        ("Subtitle", 15, 0, 8, GRAY),
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "CSE628 | Facebook Engagement and Network Visualization"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = GRAY
    _add_page_number(section.footer.paragraphs[0])


def _add_figure(
    document: Document,
    path: Path,
    number: int,
    caption: str,
    alt_text: str,
) -> None:
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.15))
    doc_properties = shape._inline.docPr
    doc_properties.set("descr", alt_text)
    caption_paragraph = document.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {number}. {caption}")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_paragraph.paragraph_format.keep_with_next = False


def _add_table(
    document: Document,
    title: str,
    headers: list[str],
    rows: Iterable[Iterable[Any]],
    widths: list[float] | None = None,
) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(9.5)
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_paragraph.paragraph_format.space_after = Pt(4)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    _set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        _set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    if widths:
        _set_table_geometry(table, widths)


def _heading_level(title: str) -> int:
    """Map numbered section depth to a Word/PDF heading level."""

    token = title.split(maxsplit=1)[0].rstrip(".")
    if token and token[0].isdigit():
        return min(token.count(".") + 1, 3)
    return 1


def build_docx(summary: dict[str, Any], path: Path) -> None:
    """Create the editable academic report with embedded tables and figures."""

    document = Document()
    _configure_document(document)
    document.add_paragraph("ACADEMIC PROJECT REPORT", style=None).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    document.add_paragraph(
        "Visual Analytics and Network Analysis of Facebook Engagement Using Python",
        style="Title",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "An empirical engagement study and seven reproducible network-visualization exercises",
        style="Subtitle",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n")
    metadata = [
        ("Student", STUDENT["name"]),
        ("Student ID", STUDENT["id"]),
        ("Course", f"{STUDENT['course_name']} ({STUDENT['course_code']})"),
        ("Semester", STUDENT["semester"]),
        ("Teacher", f"{STUDENT['teacher']}, {STUDENT['designation']}"),
        ("Department", STUDENT["department"]),
        ("University", STUDENT["university"]),
    ]
    _add_table(
        document,
        "Submission details",
        ["Field", "Information"],
        metadata,
        widths=[1.55, 4.95],
    )
    document.add_paragraph("\nPrepared: 27 July 2026").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    document.add_page_break()

    # Structured DOCX content mirrors the Markdown but uses Word-native form factors.
    sections = _docx_sections(summary)
    figure_number = 1
    for title, paragraphs, table_spec, figure_spec in sections:
        document.add_heading(title, level=_heading_level(title))
        for paragraph_text in paragraphs:
            paragraph = document.add_paragraph(paragraph_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.widow_control = True
        if table_spec:
            _add_table(document, **table_spec)
        if figure_spec:
            _add_figure(
                document,
                figure_spec[0],
                figure_number,
                figure_spec[1],
                figure_spec[2],
            )
            figure_number += 1
    document.add_heading("References", level=1)
    references = [
        "Barabasi, A.-L. (2016). Network science. Cambridge University Press. http://networksciencebook.com/",
        "Dehouche, N. (2018). Facebook Live Sellers in Thailand [Data set]. UCI Machine Learning Repository. https://doi.org/10.24432/C5R60S",
        "Dehouche, N. (2020). Dataset on usage and engagement patterns for Facebook Live sellers in Thailand. Data in Brief, 30, 105661. https://doi.org/10.1016/j.dib.2020.105661",
        "Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network structure, dynamics, and function using NetworkX. Proceedings of the 7th Python in Science Conference, 11-15.",
        "Matplotlib Development Team. (2026). Matplotlib documentation. https://matplotlib.org/stable/",
        "NetworkX Developers. (2026). NetworkX documentation. https://networkx.org/documentation/stable/",
        "pandas development team. (2026). pandas documentation. https://pandas.pydata.org/docs/",
        "Plotly Technologies Inc. (2026). Plotly Python documentation. https://plotly.com/python/",
    ]
    for reference in references:
        paragraph = document.add_paragraph(reference)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    document.add_heading("Appendix: Reproducibility and assumptions", level=1)
    appendix_paragraph = document.add_paragraph(
        "Run python main.py to rebuild every generated artifact and pytest to validate "
        "the pipeline. Each exercise folder contains its standalone Python source, clean "
        "notebook, executed notebook, and browser-ready HTML notebook. The supplied teaching "
        "notebooks are source-identical. The official UCI download was used because Kaggle "
        "credentials were unavailable. All assignment-only relationship graphs are labeled "
        "synthetic."
    )
    appendix_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.core_properties.title = (
        "Visual Analytics and Network Analysis of Facebook Engagement Using Python"
    )
    document.core_properties.author = STUDENT["name"]
    document.core_properties.subject = f"{STUDENT['course_code']} academic project"
    document.save(path)
    _normalize_docx_container(path)


def _exercise_docx_sections(
    summary: dict[str, Any],
) -> list[tuple[str, list[str], dict[str, Any] | None, tuple[Path, str, str] | None]]:
    """Return detailed Word/PDF sections for the seven exercises."""

    exercises = summary["exercises"]
    transport = exercises["exercise_1"]
    ppi = exercises["exercise_2"]
    bipartite = exercises["exercise_3"]
    weighted = exercises["exercise_4"]
    models = exercises["exercise_5"]
    social = exercises["exercise_6"]
    domain = exercises["exercise_7"]
    matrix_rows = _transport_matrix_rows(transport)
    model_rows = [
        [
            row["Model"],
            row["Nodes"],
            row["Edges"],
            f"{row['Density']:.3f}",
            f"{row['Average Clustering']:.3f}",
            f"{row['Average Shortest Path']:.3f}",
            row["Connected Components"],
            row["Maximum Degree"],
        ]
        for row in models["records"]
    ]
    course_rows = [
        [row["course"], row["enrollments"]] for row in bipartite["course_degrees"]
    ]
    weight_rows = [
        [weight, count] for weight, count in weighted["weight_frequencies"].items()
    ]
    group_rows = [
        [group, count] for group, count in social["interest_group_counts"].items()
    ]
    centrality_rows = [
        [
            row["label"],
            row["node_type"],
            row["degree"],
            f"{row['betweenness_centrality']:.4f}",
            f"{row['closeness_centrality']:.4f}",
            f"{row['pagerank']:.4f}",
        ]
        for row in domain["top_centrality_records"]
    ]
    top_users = ", ".join(
        f"{row['user']} ({row['in_degree']})"
        for row in social["highest_in_degree_users"]
    )
    sections = [
        (
            "7. Exercise 1 - Weighted adjacency representations",
            [
                "The relationship between an adjacency list and a weighted adjacency "
                "matrix was examined through the supplied G_transport network. Matrix "
                "symmetry and its implication for route direction were evaluated."
            ],
            None,
            None,
        ),
        (
            "7.1 Objective, construction, and procedure",
            [
                (
                    f"Seven cities and {transport['edges']} undirected teaching routes "
                    "were represented. Distance in kilometers was stored in every edge "
                    f"weight. Density was {transport['density']:.4f}, and Dhaka received "
                    f"the largest degree ({transport['highest_degree']})."
                ),
                (
                    "A labeled adjacency list was produced from neighbor iterators. A "
                    "weighted matrix was generated with nx.to_pandas_adjacency, and "
                    "symmetry was tested by comparison with the transpose through "
                    "numpy.allclose. Off-diagonal zeros were treated as absent routes."
                ),
            ],
            None,
            None,
        ),
        (
            "7.2 Verified results",
            [
                (
                    f"The symmetry test returned {transport['matrix_symmetric']}. The "
                    f"{transport['longest_route']['source']}-"
                    f"{transport['longest_route']['target']} edge was longest at "
                    f"{transport['longest_route']['distance_km']} km, while the "
                    f"{transport['shortest_route']['source']}-"
                    f"{transport['shortest_route']['target']} edge was shortest at "
                    f"{transport['shortest_route']['distance_km']} km. Unique edge "
                    f"distances totaled {transport['total_recorded_distance_km']:,} km."
                )
            ],
            {
                "title": "Table 2. Weighted adjacency matrix (kilometers)",
                "headers": ["City", *transport["node_order"]],
                "rows": matrix_rows,
                "widths": [1.0, *([0.7854166667] * 6), 0.7875],
            },
            None,
        ),
        (
            "7.3 Interpretation, limitations, and verification",
            [
                (
                    "Symmetry was produced because every route was encoded as an "
                    "undirected edge with one shared weight. An asymmetric matrix would "
                    "be expected under one-way travel or direction-dependent cost. The "
                    "adjacency list was better suited to neighbor inspection, while the "
                    "matrix supported exact pairwise lookup."
                ),
                (
                    "The graph was treated as a classroom illustration rather than a "
                    "complete transportation model. The matrix was saved as a CSV, and "
                    "symmetry is covered by an automated graph-representation test."
                ),
            ],
            None,
            None,
        ),
        (
            "8. Exercise 2 - Six-layout comparison",
            [
                "Six layouts were compared for one unchanged G_ppi topology. Identity, "
                "labels, node size, color, and edges were fixed so that placement was "
                "the only changing visual variable."
            ],
            None,
            None,
        ),
        (
            "8.1 Graph construction and controlled procedure",
            [
                (
                    "The graph was reproduced with nx.watts_strogatz_graph using n=15, "
                    "k=3, p=0.3, and seed 42, then relabeled P1 through P15. The "
                    f"realization contains {ppi['edges']} edges, density "
                    f"{ppi['density']:.4f}, average degree {ppi['average_degree']:.2f}, "
                    f"diameter {ppi['diameter']}, and degree range "
                    f"{ppi['minimum_degree']}-{ppi['maximum_degree']}."
                ),
                (
                    "Spring, circular, shell, spectral, Kamada-Kawai, and random "
                    "placement were assessed through crossings, label separation, ring "
                    "locality, shortcut visibility, and correspondence between graph "
                    "distance and visual proximity."
                ),
            ],
            None,
            None,
        ),
        (
            "8.2 Verified result and critical interpretation",
            [
                (
                    "Kamada-Kawai was selected for this realization because pairwise "
                    "distance stress was minimized and ring-like locality remained "
                    "legible. Spring placement was a close alternative."
                ),
                (
                    f"Average clustering was {ppi['average_clustering']:.3f}. "
                    "Triangle-based clustering was absent in this exact seeded graph, "
                    "so the layouts were interpreted as views of locality and shortcuts "
                    "rather than as evidence of communities."
                ),
            ],
            None,
            (
                STATIC / "g_ppi_layout_comparison.png",
                "Six layouts of the identical synthetic G_ppi graph.",
                "Spring, circular, shell, spectral, Kamada-Kawai, and random layouts of one graph.",
            ),
        ),
        (
            "8.3 Limitations and verification",
            [
                "Layout preference remains dependent on graph size, parameters, task, "
                "and labeling. Six individual figures and a combined comparison were "
                "saved, and all reported graph statistics were recalculated from the "
                "generated NetworkX object."
            ],
            None,
            None,
        ),
        (
            "9. Exercise 3 - Student-course bipartite network",
            [
                "A two-mode enrollment network was constructed without inventing "
                "student-student or course-course edges."
            ],
            None,
            None,
        ),
        (
            "9.1 Data design, construction procedure, and validation",
            [
                (
                    f"The stored synthetic data contain {bipartite['students']} students, "
                    f"{bipartite['courses']} courses, and {bipartite['enrollments']} "
                    "enrollment records. Students were assigned to partition 0 and "
                    "courses to partition 1."
                ),
                (
                    "An edge was added only when a stored enrollment linked the two "
                    f"partitions. Bipartiteness returned {bipartite['is_bipartite']}, "
                    f"and bipartite density was {bipartite['bipartite_density']:.4f}."
                ),
            ],
            None,
            None,
        ),
        (
            "9.2 Degree results",
            [
                (
                    f"Mean course load was {bipartite['mean_courses_per_student']:.2f} "
                    "per student, and mean course enrollment was "
                    f"{bipartite['mean_students_per_course']:.2f}. "
                    f"{bipartite['most_popular_course']} had the largest course degree "
                    f"({bipartite['most_popular_course_enrollments']}), while "
                    f"{bipartite['most_enrolled_student']} had the largest student "
                    f"degree ({bipartite['most_enrolled_student_courses']})."
                )
            ],
            {
                "title": "Table 3. Enrollment degree by course",
                "headers": ["Course", "Student enrollments"],
                "rows": course_rows,
                "widths": [4.8, 1.7],
            },
            None,
        ),
        (
            "9.3 Visual result and interpretation",
            [
                (
                    "Separate columns, shapes, and colors made partition membership "
                    "explicit. Student degree represented course load, whereas course "
                    "degree represented synthetic popularity. A one-mode projection was "
                    "not used because derived ties were outside the exercise scope."
                ),
                bipartite["pattern"],
            ],
            None,
            (
                STATIC / "student_course_bipartite_graph.png",
                "Synthetic student-course enrollment network.",
                "Bipartite graph with student circles and course squares.",
            ),
        ),
        (
            "9.4 Limitations and verification",
            [
                "Names and enrollment records were generated solely for teaching. No "
                "claim was made about actual students or enrollment behavior. Source "
                "tables, calculated degree results, and an automated partition test "
                "provide reproduction evidence."
            ],
            None,
            None,
        ),
        (
            "10. Exercise 4 - Weighted Barabasi-Albert graph",
            [
                "The perceptual effect of edge-width encoding was examined while the "
                "underlying 100-node preferential-attachment topology was preserved."
            ],
            None,
            None,
        ),
        (
            "10.1 Construction and weight assignment",
            [
                (
                    "The topology was generated with n=100, m=2, and seed 42. It "
                    f"contains {weighted['edges']} edges, average degree "
                    f"{weighted['average_degree']:.2f}, and maximum degree "
                    f"{weighted['maximum_degree']}."
                ),
                (
                    "Each copied edge received a reproducible integer weight from 1 "
                    "through 10. Edge width was scaled from weight; topology, degree, "
                    "and node positions were not changed."
                ),
            ],
            None,
            None,
        ),
        (
            "10.2 Verified weight distribution",
            [
                (
                    f"The observed range was {weighted['minimum_weight']}-"
                    f"{weighted['maximum_weight']}, the mean was "
                    f"{weighted['mean_weight']:.3f}, the median was "
                    f"{weighted['median_weight']:.1f}, and the population standard "
                    f"deviation was {weighted['weight_standard_deviation']:.3f}. Total "
                    f"assigned edge weight was {weighted['total_edge_weight']:,}."
                )
            ],
            {
                "title": "Table 4. Frequency of assigned edge weights",
                "headers": ["Edge weight", "Number of edges"],
                "rows": weight_rows,
                "widths": [2.4, 4.1],
            },
            (
                STATIC / "barabasi_albert_weighted.png",
                "Barabasi-Albert graph with edge-width encoding.",
                "Weighted synthetic preferential-attachment graph with thicker high-weight edges.",
            ),
        ),
        (
            "10.3 Interpretation, limitations, and verification",
            [
                "Thicker lines increased tie salience while hub structure remained "
                "unchanged. Edge weight was not treated as node centrality. A weighted "
                "centrality analysis would require weight to be defined as strength, "
                "capacity, cost, or distance.",
                "The random weights demonstrate encoding only. The edge table was saved, "
                "and automated tests verify graph order, size, range, and determinism.",
            ],
            None,
            None,
        ),
        (
            "11. Exercise 5 - Generative-model comparison",
            [
                "Random mixing, local clustering with rewiring, and preferential "
                "attachment were compared under the exact supplied parameters."
            ],
            None,
            None,
        ),
        (
            "11.1 Model construction, parameters, and metric procedure",
            [
                "Erdos-Renyi G(30, 0.08), Watts-Strogatz WS(30, 4, 0.1), and "
                "Barabasi-Albert BA(30, 2) were generated with seed 42.",
                "Order, size, density, average degree, clustering, connectivity, "
                "components, average path length, maximum degree, and degree standard "
                "deviation were calculated. For the disconnected Erdos-Renyi graph, "
                "path length was calculated on the largest connected component.",
            ],
            None,
            None,
        ),
        (
            "11.2 Verified statistical results",
            [],
            {
                "title": "Table 5. Statistical comparison of seeded models",
                "headers": [
                    "Model",
                    "Nodes",
                    "Edges",
                    "Density",
                    "Clustering",
                    "Avg. path",
                    "Components",
                    "Max degree",
                ],
                "rows": model_rows,
                "widths": [1.65, 0.5, 0.5, 0.65, 0.7, 0.7, 0.8, 1.0],
            },
            (
                STATIC / "generative_models_degree_distribution.png",
                "Degree distributions of the three seeded models.",
                "Degree histograms for Erdos-Renyi, Watts-Strogatz, and Barabasi-Albert graphs.",
            ),
        ),
        (
            "11.3 Model-by-model interpretation",
            [
                (
                    "The Erdos-Renyi realization produced three components and low "
                    f"clustering ({models['records'][0]['Average Clustering']:.3f}); it "
                    "served as a homogeneous random baseline."
                ),
                (
                    "The Watts-Strogatz realization remained connected and produced the "
                    f"highest clustering ({models['records'][1]['Average Clustering']:.3f}) "
                    "with a narrow degree distribution."
                ),
                (
                    "The Barabasi-Albert realization remained connected, produced the "
                    f"shortest average path ({models['records'][2]['Average Shortest Path']:.3f}), "
                    f"and generated maximum degree {models['records'][2]['Maximum Degree']} "
                    "with strong degree heterogeneity."
                ),
                models["conclusion"],
            ],
            None,
            None,
        ),
        (
            "11.4 Limitations and verification",
            [
                "Only one small seeded realization was compared per model. Repeated "
                "simulation and uncertainty intervals would be required for general "
                "model inference. The comparison CSV and separate degree-distribution "
                "figures provide reproduction evidence."
            ],
            None,
            None,
        ),
        (
            "12. Exercise 6 - Interactive node-color dashboard",
            [
                "One directed graph was interpreted through categorical interest-group "
                "color and continuous in-degree color without changing node positions."
            ],
            None,
            None,
        ),
        (
            "12.1 Graph construction procedure and results",
            [
                (
                    f"The seeded logic produced {social['nodes']} users and "
                    f"{social['edges']} directed follow edges. Directed density was "
                    f"{social['density']:.4f}. Mean in-degree was "
                    f"{social['mean_in_degree']:.2f}; in-degree ranged from "
                    f"{social['minimum_in_degree']} to {social['maximum_in_degree']}, "
                    f"and out-degree ranged from {social['minimum_out_degree']} to "
                    f"{social['maximum_out_degree']}."
                ),
                f"The three largest in-degree results were {top_users}.",
            ],
            {
                "title": "Table 6. Synthetic interest-group membership",
                "headers": ["Interest group", "Users"],
                "rows": group_rows,
                "widths": [4.8, 1.7],
            },
            None,
        ),
        (
            "12.2 Initial categorical-color state",
            [
                "Categorical color was used to support comparison of synthetic interest "
                "groups. Node coordinates, edges, labels, and hover fields were fixed."
            ],
            None,
            (
                REPORT.parent / "screenshots" / "network-toggle-before.png",
                "Interactive network colored by interest group.",
                "Network dashboard in the categorical interest-group color mode.",
            ),
        ),
        (
            "12.3 Continuous in-degree state",
            [
                (
                    "The dropdown changed marker color, legend, and scale metadata to "
                    "in-degree while coordinates remained fixed. Position preservation "
                    f"was verified as {social['positions_preserved']}."
                )
            ],
            None,
            (
                REPORT.parent / "screenshots" / "network-toggle-in-degree.png",
                "Interactive network colored by in-degree.",
                "Network dashboard in the continuous in-degree color mode.",
            ),
        ),
        (
            "12.4 Interpretation, limitations, and verification",
            [
                "The categorical state supported group-mixing interpretation, while the "
                "continuous state emphasized incoming-tie popularity. Perceptual change "
                "was attributable to color because position was preserved.",
                "The graph is synthetic, and Plotly line traces do not display arrowheads. "
                "Direction remains available through degree metrics. Both dropdown states "
                "were verified through browser interaction.",
            ],
            None,
            None,
        ),
        (
            "13. Exercise 7 - Research-domain knowledge graph",
            [
                "A typed synthetic graph was used to represent research areas, methods, "
                "tools, applications, outcomes, and semantic relationships."
            ],
            None,
            None,
        ),
        (
            "13.1 Construction and metric procedure",
            [
                (
                    f"The graph contains {domain['nodes']} nodes and "
                    f"{domain['edges']} undirected links. Density was "
                    f"{domain['density']:.4f}, average degree was "
                    f"{domain['average_degree']:.2f}, and connectivity was verified as "
                    f"{domain['connected']}."
                ),
                "Degree centrality, inverse-strength weighted betweenness, closeness, "
                "and strength-weighted PageRank were calculated. Edge distance was "
                "defined as 1 / strength before shortest-path betweenness was evaluated.",
            ],
            None,
            None,
        ),
        (
            "13.2 Verified centrality results",
            [
                (
                    f"{domain['top_betweenness_node']} received the largest weighted "
                    f"betweenness value ({domain['top_betweenness_value']:.4f})."
                )
            ],
            {
                "title": "Table 7. Five highest weighted-betweenness nodes",
                "headers": [
                    "Node",
                    "Type",
                    "Degree",
                    "Betweenness",
                    "Closeness",
                    "PageRank",
                ],
                "rows": centrality_rows,
                "widths": [1.45, 1.25, 0.6, 1.05, 1.05, 1.1],
            },
            (
                STATIC / "domain_graph.png",
                "Synthetic Applied AI and Multimedia knowledge graph.",
                "Typed knowledge graph with weighted semantic relationships.",
            ),
        ),
        (
            "13.3 Interpretation, limitations, and verification",
            [
                domain["interpretation"],
                (
                    "The graph was not extracted from publications. The inverse-strength "
                    "transformation assumes that stronger semantic relationships represent "
                    "shorter effective distances. A different substantive meaning for "
                    "strength would require a different transformation."
                ),
                "Stored node and edge tables, calculated centrality metrics, and static "
                "and interactive figures provide reproduction evidence.",
            ],
            None,
            None,
        ),
    ]
    normalized = []
    exercise_number = 0
    subsection_number = 0
    for title, paragraphs, table_spec, figure_spec in sections:
        if "Exercise " in title:
            exercise_number += 1
            subsection_number = 0
            label = title.split(" - ", maxsplit=1)[-1]
            normalized_title = (
                f"8.{exercise_number} Exercise {exercise_number} - {label}"
            )
        else:
            subsection_number += 1
            label = title.split(maxsplit=1)[-1]
            normalized_title = f"8.{exercise_number}.{subsection_number} {label}"
        normalized.append((normalized_title, paragraphs, table_spec, figure_spec))
    return normalized


def _docx_sections(
    summary: dict[str, Any],
) -> list[tuple[str, list[str], dict[str, Any] | None, tuple[Path, str, str] | None]]:
    engagement = summary["engagement"]
    audit = summary["data_cleaning"]
    exercises = summary["exercises"]
    model_records = exercises["exercise_5"]["records"]
    type_rows = [
        [
            row["status_type"].title(),
            f"{row['posts']:,}",
            f"{row['mean_reactions']:,.2f}",
            f"{row['mean_comments']:,.2f}",
            f"{row['mean_shares']:,.2f}",
            f"{row['mean_total_engagement']:,.2f}",
        ]
        for row in engagement["by_type"]
    ]
    model_rows = [
        [
            row["Model"],
            row["Nodes"],
            row["Edges"],
            f"{row['Density']:.3f}",
            f"{row['Average Clustering']:.3f}",
            f"{row['Average Shortest Path']:.3f}",
            row["Maximum Degree"],
        ]
        for row in model_records
    ]
    sections = [
        (
            "Abstract",
            [
                (
                    f"An analysis is presented of {engagement['rows']:,} anonymized "
                    "Facebook posts, and seven reproducible network-visualization "
                    "exercises are completed. Cleaning, feature engineering, static and "
                    "interactive visualization, graph statistics, testing, and public "
                    "web delivery are combined. The largest observed mean total "
                    "engagement is associated with video posts "
                    f"({engagement['top_post_type_mean_engagement']:,.2f}), but causality "
                    "is not established by the observational design."
                ),
                (
                    "Keywords: Facebook engagement, data visualization, social-network "
                    "analysis, NetworkX, Plotly, Python."
                ),
            ],
            None,
            None,
        ),
        (
            "Contents",
            [
                "1. Introduction and research questions; 2. Dataset, license, and "
                "ethics; 3. Data cleaning and feature engineering; 4-6. Empirical "
                "findings and visualization methodology; 7. Extended analytical "
                "framework; 8. Seven mandatory exercises; 9-12. Discussion, "
                "limitations, conclusions, implications, references, and appendices."
            ],
            None,
            None,
        ),
        (
            "1. Introduction and research questions",
            [
                "Multidimensional engagement is described while the distinction between "
                "real tabular observations and synthetic relationship graphs is "
                "preserved. Variation by content and time, structural readability under "
                "different visual layouts, and the social-network properties captured by "
                "canonical graph models are examined.",
                "A central methodological constraint is that the Facebook table contains "
                "no verified user-to-user edges. The empirical analysis therefore remains "
                "tabular; all assignment relationship graphs are explicitly synthetic.",
            ],
            None,
            None,
        ),
        (
            "2. Dataset, license, and ethics",
            [
                (
                    "Facebook Live Sellers in Thailand was authored by Nassim Dehouche "
                    "and distributed by UCI under CC BY 4.0. It records anonymized post "
                    f"engagement from {engagement['date_min'][:10]} through "
                    f"{engagement['date_max'][:10]}."
                ),
                "The official UCI archive was used because Kaggle credentials were not "
                "configured. The Kaggle mirror is documented but was not the acquisition "
                "route. No names, message text, or re-identification attempts are included.",
            ],
            None,
            None,
        ),
        (
            "3. Data cleaning and feature engineering",
            [
                (
                    f"The raw table contained {audit['initial_rows']:,} rows and "
                    f"{audit['initial_columns']} columns. Four empty placeholder columns "
                    f"were removed; {audit['duplicate_rows_removed']} duplicate rows and "
                    f"{audit['invalid_dates_removed']} invalid dates were found. The "
                    f"processed table contains {audit['final_rows']:,} rows and "
                    f"{audit['final_columns']} columns."
                ),
                (
                    "Total engagement equals reactions plus comments plus shares. Ratio "
                    "features use zero-safe division. Temporal features include posting "
                    "hour, weekday, month, ISO week, weekend, and time-of-day. No engagement "
                    "rate is claimed because reach and impressions are absent."
                ),
            ],
            None,
            None,
        ),
        (
            "4. Exploratory engagement findings",
            [
                (
                    f"The data record {engagement['total_reactions']:,} reactions, "
                    f"{engagement['total_comments']:,} comments, and "
                    f"{engagement['total_shares']:,} shares. Mean total engagement is "
                    f"{engagement['mean_engagement']:,.2f}, while the median is "
                    f"{engagement['median_engagement']:,.2f}; the gap is consistent with "
                    "strong right skew."
                )
            ],
            {
                "title": "Table 1. Engagement summary by post type",
                "headers": [
                    "Type",
                    "Posts",
                    "Mean reactions",
                    "Mean comments",
                    "Mean shares",
                    "Mean total",
                ],
                "rows": type_rows,
                "widths": [0.7, 0.65, 1.15, 1.15, 1.05, 1.8],
            },
            (
                STATIC / "engagement_by_post_type.png",
                "Mean engagement components by post type.",
                "Grouped bar chart of mean reactions, comments, and shares by post type.",
            ),
        ),
        (
            "5. Distribution, correlation, and time",
            [
                (
                    f"The upper Tukey fence flags {audit['outlier_count']:,} posts rather "
                    "than deleting them. Rank correlations and log-transformed distributions "
                    "are used because extreme posts make raw-count views difficult to read."
                ),
                (
                    f"The largest observed median engagement occurs at "
                    f"{engagement['best_median_posting_hour']:02d}:00, but content mix, "
                    "seller behavior, and seasonality are plausible confounders."
                ),
            ],
            None,
            (
                STATIC / "engagement_distribution.png",
                "Log-transformed total-engagement distributions by post type.",
                "Histogram comparison of log transformed engagement distributions.",
            ),
        ),
        (
            "5.1 Rank-correlation result",
            [
                "Spearman rank correlation was used because the engagement variables "
                "are strongly skewed. The coefficients describe monotonic association "
                "and were not interpreted as causal effects."
            ],
            None,
            (
                STATIC / "engagement_correlation.png",
                "Spearman correlation matrix for engagement measures.",
                "Heatmap of rank correlations among reactions, comments, shares, and total engagement.",
            ),
        ),
        (
            "6. Graph representation and visualization methodology",
            [
                "Adjacency matrices support exact pairwise lookup; node-link diagrams "
                "emphasize paths, clusters, bridges, and hubs. Visual encodings are held "
                "constant during layout comparison so that topology-driven placement is "
                "the primary changing variable.",
                "Force-directed algorithms map graph-theoretic relationships to proximity, "
                "whereas circular and random layouts do not optimize community readability.",
            ],
            None,
            None,
        ),
        (
            "7. Exercise 1 - Weighted adjacency representations",
            [
                (
                    "The G_transport matrix stores route distances in kilometers and is "
                    f"symmetric: {exercises['exercise_1']['matrix_symmetric']}. Symmetry "
                    "follows from undirected equal-weight edges; a directed or directionally "
                    "weighted network could be asymmetric."
                )
            ],
            None,
            None,
        ),
        (
            "8. Exercise 2 - Six-layout comparison",
            [
                "Kamada-Kawai is selected as the clearest layout for this exact 15-node "
                "G_ppi realization because it minimizes pairwise distance stress and "
                "makes ring locality plus rewired shortcuts legible. The measured average "
                f"clustering is {exercises['exercise_2']['average_clustering']:.3f}; no "
                "layout can reveal triangle-based clustering that is absent. Spring is a "
                "close alternative, and the judgment is not universal."
            ],
            None,
            (
                STATIC / "g_ppi_layout_comparison.png",
                "Six layouts of the identical synthetic G_ppi graph.",
                "Comparison of spring, circular, shell, spectral, Kamada-Kawai, and random layouts.",
            ),
        ),
        (
            "9. Exercise 3 - Student-course bipartite network",
            [
                (
                    f"The stored CSV tables form a valid bipartite graph with "
                    f"{exercises['exercise_3']['enrollments']} enrollments. "
                    f"{exercises['exercise_3']['most_popular_course']} is most popular, "
                    f"and {exercises['exercise_3']['most_enrolled_student']} takes the "
                    "largest number of courses."
                ),
                exercises["exercise_3"]["pattern"],
            ],
            None,
            (
                STATIC / "student_course_bipartite_graph.png",
                "Synthetic student-course enrollment network.",
                "Bipartite graph with student circles and course squares.",
            ),
        ),
        (
            "10. Exercise 4 - Weighted Barabasi-Albert graph",
            [
                (
                    f"All {exercises['exercise_4']['edges']} edges received seeded integer "
                    f"weights from {exercises['exercise_4']['minimum_weight']} to "
                    f"{exercises['exercise_4']['maximum_weight']}. Width changes tie "
                    "salience but does not redefine degree-based hubs. Random weights "
                    "demonstrate encoding rather than empirical importance."
                )
            ],
            None,
            (
                STATIC / "barabasi_albert_weighted.png",
                "Barabasi-Albert graph with normalized edge-width encoding.",
                "Weighted synthetic scale-free graph with thicker high-weight edges.",
            ),
        ),
        (
            "11. Exercise 5 - Generative-model comparison",
            [
                "Watts-Strogatz best captures high clustering and short paths; "
                "Barabasi-Albert best captures hubs and heterogeneous degree. A realistic "
                "social network may combine both, so none is universally best. "
                "Erdos-Renyi remains a useful homogeneous random baseline."
            ],
            {
                "title": "Table 2. Statistical comparison of Section 3 models",
                "headers": [
                    "Model",
                    "Nodes",
                    "Edges",
                    "Density",
                    "Clustering",
                    "Avg. path",
                    "Max degree",
                ],
                "rows": model_rows,
                "widths": [1.65, 0.55, 0.55, 0.7, 0.75, 0.75, 0.75],
            },
            (
                STATIC / "generative_models_degree_distribution.png",
                "Degree distributions of the three Section 3 models.",
                "Side-by-side degree histograms for three synthetic graph models.",
            ),
        ),
        (
            "12. Exercise 6 - Interactive node-color dashboard",
            [
                "The Plotly dashboard preserves node positions while a dropdown switches "
                "between categorical interest-group color and continuous in-degree color. "
                "Hover information reports node ID, group, in-degree, and out-degree, and "
                "edges remain visible in both modes."
            ],
            None,
            None,
        ),
        (
            "13. Exercise 7 - Applied AI and Multimedia research graph",
            [
                exercises["exercise_7"]["interpretation"],
                "The graph is synthetic and demonstrates centrality interpretation rather "
                "than empirical evidence about a research community.",
            ],
            None,
            (
                STATIC / "domain_graph.png",
                "Synthetic applied-AI and multimedia research knowledge graph.",
                "Typed research knowledge graph with weighted semantic relationships.",
            ),
        ),
        (
            "9. Discussion",
            [
                "The empirical and synthetic components answer different questions. The "
                "Facebook table supports descriptive post comparison but not individual "
                "relationship inference. The synthetic graphs reveal how topology, layout, "
                "and encoding shape network interpretation.",
                "Mean engagement is influenced by extreme posts, so median summaries, log "
                "distributions, and outlier flags accompany it. Correlations remain "
                "associational and do not isolate algorithm, audience, or content effects.",
            ],
            None,
            None,
        ),
        (
            "10. Limitations and ethical considerations",
            [
                "The data end in 2018, represent ten Thai sellers, and omit reach, "
                "impressions, followers, campaign spend, page identity, and content text. "
                "Results may not generalize across settings or time.",
                "The project preserves anonymization, reports aggregate findings, avoids "
                "re-identification, makes no unsupported causal claims, and clearly labels "
                "every dummy graph as synthetic.",
            ],
            None,
            None,
        ),
        (
            "11. Conclusion, recommendations, and future work",
            [
                "The largest observed mean engagement was associated with video posts. "
                "The supplied small-world graph was presented most clearly by "
                "Kamada-Kawai, tie salience was separated from centrality through weighted "
                "encoding, and the multidimensional nature of social-network realism was "
                "revealed by the model comparison.",
                "For future work, ethically collected reach denominators, multiple seller "
                "populations, repeated graph simulations, uncertainty intervals, content "
                "representations, and user testing of dashboard accessibility should be "
                "added.",
            ],
            None,
            None,
        ),
    ]
    first_exercise = next(
        index
        for index, section in enumerate(sections)
        if section[0].startswith("7. Exercise 1")
    )
    discussion = next(
        index
        for index, section in enumerate(sections)
        if section[0].startswith("9. Discussion")
    )
    extended = _extended_academic_sections(summary)
    pre_exercise = [section for section in extended if section[0].startswith("7.")]
    post_exercise = [section for section in extended if section[0].startswith("12.")]
    return (
        sections[:first_exercise]
        + pre_exercise
        + _exercise_docx_sections(summary)
        + sections[discussion:]
        + post_exercise
    )


def _configure_pdf_styles():
    """Return report styles with justified academic body paragraphs."""

    styles = getSampleStyleSheet()
    styles["Heading1"].keepWithNext = 1
    styles["Heading2"].keepWithNext = 1
    styles["Heading3"].keepWithNext = 1
    styles.add(
        ParagraphStyle(
            name="AcademicBody",
            parent=styles["BodyText"],
            alignment=TA_JUSTIFY,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=29,
            textColor=colors.HexColor("#17324D"),
            alignment=TA_CENTER,
            spaceAfter=18,
        )
    )
    return styles


def build_pdf(summary: dict[str, Any], path: Path) -> None:
    """Create a compact PDF fallback from the same verified summary."""

    styles = _configure_pdf_styles()
    document = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title="Visual Analytics and Network Analysis of Facebook Engagement Using Python",
        author=STUDENT["name"],
    )
    story = [
        Spacer(1, 1.2 * inch),
        Paragraph(
            "Visual Analytics and Network Analysis of Facebook Engagement Using Python",
            styles["CoverTitle"],
        ),
        Paragraph(
            f"{STUDENT['name']} | {STUDENT['id']}<br/>"
            f"{STUDENT['course_name']} ({STUDENT['course_code']}) | "
            f"{STUDENT['semester']}<br/>"
            f"{STUDENT['teacher']}, {STUDENT['designation']}<br/>"
            f"{STUDENT['department']}<br/>{STUDENT['university']}",
            styles["BodyText"],
        ),
        PageBreak(),
    ]
    figure_number = 1
    for title, paragraphs, table_spec, figure_spec in _docx_sections(summary):
        heading = Paragraph(title, styles[f"Heading{_heading_level(title)}"])
        remaining_paragraphs = paragraphs
        if paragraphs:
            story.append(heading)
            story.append(Paragraph(paragraphs[0], styles["AcademicBody"]))
            story.append(Spacer(1, 0.08 * inch))
            remaining_paragraphs = paragraphs[1:]
        else:
            story.append(heading)
        for paragraph in remaining_paragraphs:
            story.append(
                KeepTogether(
                    [
                        Paragraph(paragraph, styles["AcademicBody"]),
                        Spacer(1, 0.08 * inch),
                    ]
                )
            )
        if table_spec:
            rows = [table_spec["headers"], *table_spec["rows"]]
            table = Table(
                rows,
                colWidths=[width * inch for width in table_spec["widths"]],
                repeatRows=1,
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E9EEF3")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 7),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB4BD")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.12 * inch))
        if figure_spec and figure_spec[0].exists():
            story.append(
                KeepTogether(
                    [
                        Image(
                            str(figure_spec[0]),
                            width=6.3 * inch,
                            height=3.9 * inch,
                        ),
                        Paragraph(
                            f"Figure {figure_number}. {figure_spec[1]}",
                            styles["Italic"],
                        ),
                    ]
                )
            )
            figure_number += 1
    story.append(Paragraph("References", styles["Heading1"]))
    references = [
        "Barabasi, A.-L. (2016). Network science. Cambridge University Press. http://networksciencebook.com/",
        "Dehouche, N. (2018). Facebook Live Sellers in Thailand [Data set]. UCI Machine Learning Repository. https://doi.org/10.24432/C5R60S",
        "Dehouche, N. (2020). Dataset on usage and engagement patterns for Facebook Live sellers in Thailand. Data in Brief, 30, 105661. https://doi.org/10.1016/j.dib.2020.105661",
        "Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network structure, dynamics, and function using NetworkX. Proceedings of the 7th Python in Science Conference, 11-15.",
        "Matplotlib Development Team. (2026). Matplotlib documentation. https://matplotlib.org/stable/",
        "NetworkX Developers. (2026). NetworkX documentation. https://networkx.org/documentation/stable/",
        "pandas development team. (2026). pandas documentation. https://pandas.pydata.org/docs/",
        "Plotly Technologies Inc. (2026). Plotly Python documentation. https://plotly.com/python/",
    ]
    for reference in references:
        story.append(Paragraph(reference, styles["BodyText"]))
        story.append(Spacer(1, 0.06 * inch))
    story.append(
        Paragraph("Appendix: Reproducibility and assumptions", styles["Heading1"])
    )
    story.append(
        Paragraph(
            "The generated artifacts can be rebuilt by running python main.py, and the "
            "pipeline can be validated with pytest. Every numbered exercise folder contains "
            "a standalone Python source, a clean notebook, a clean-kernel executed notebook, "
            "and a browser-ready HTML notebook; the interactive exercises also retain their "
            "Plotly HTML. The supplied teaching notebooks are source-identical. The official "
            "UCI archive was used because Kaggle credentials were unavailable. All assignment "
            "relationship graphs are explicitly synthetic.",
            styles["AcademicBody"],
        )
    )

    def add_page_furniture(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#52606D"))
        canvas.drawString(
            0.75 * inch,
            0.42 * inch,
            "CSE628 | Facebook Engagement and Network Visualization",
        )
        canvas.drawRightString(
            letter[0] - 0.75 * inch,
            0.42 * inch,
            f"Page {doc.page}",
        )
        canvas.restoreState()

    document.build(
        story,
        onFirstPage=add_page_furniture,
        onLaterPages=add_page_furniture,
    )


def generate_reports(summary_path: Path = ANALYSIS_SUMMARY) -> dict[str, Path]:
    """Generate all requested report formats."""

    summary = _load_summary(summary_path)
    REPORT.mkdir(parents=True, exist_ok=True)
    markdown_path = REPORT / "report.md"
    docx_path = REPORT / "report.docx"
    pdf_path = REPORT / "report.pdf"
    markdown_path.write_text(build_markdown(summary), encoding="utf-8")
    build_docx(summary, docx_path)
    build_pdf(summary, pdf_path)
    return {"markdown": markdown_path, "docx": docx_path, "pdf": pdf_path}
