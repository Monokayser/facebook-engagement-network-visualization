"""Academic-structure and voice checks for the generated report."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
from reportlab.lib.enums import TA_JUSTIFY

from src.report_generator import _configure_pdf_styles

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
FIRST_PERSON = re.compile(r"\b(?:I|me|my|mine|we|us|our|ours)\b", re.IGNORECASE)


def _docx_text() -> str:
    document = Document(REPORT / "report.docx")
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *table_cells])


def _pdf_text() -> str:
    return "\n".join(
        page.extract_text() or "" for page in PdfReader(REPORT / "report.pdf").pages
    )


def test_each_exercise_has_individual_academic_subsections() -> None:
    markdown = (REPORT / "report.md").read_text(encoding="utf-8")
    required_parts = [
        "Objective",
        "construction",
        "procedure",
        "Verified",
        "interpretation",
        "Limitations",
        "verification",
    ]
    for exercise in range(1, 8):
        start = markdown.index(f"### 8.{exercise} Exercise {exercise}")
        end_marker = (
            f"### 8.{exercise + 1} Exercise {exercise + 1}"
            if exercise < 7
            else "## 9. Discussion"
        )
        section = markdown[start : markdown.index(end_marker, start)]
        for part in required_parts:
            assert part.lower() in section.lower(), (exercise, part)


def test_report_formats_exclude_first_person_pronouns() -> None:
    markdown = (REPORT / "report.md").read_text(encoding="utf-8")
    for text in (markdown, _docx_text(), _pdf_text()):
        assert FIRST_PERSON.search(text) is None


def test_report_formats_include_all_seven_exercises() -> None:
    for text in (
        (REPORT / "report.md").read_text(encoding="utf-8"),
        _docx_text(),
        _pdf_text(),
    ):
        for exercise in range(1, 8):
            assert f"Exercise {exercise}" in text


def test_academic_body_paragraphs_are_justified() -> None:
    document = Document(REPORT / "report.docx")
    normal = document.styles["Normal"]
    assert normal.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    body_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "Normal"
        and len(paragraph.text.split()) >= 20
        and paragraph.paragraph_format.left_indent is None
    ]
    assert body_paragraphs
    assert all(
        paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.JUSTIFY)
        for paragraph in body_paragraphs
    )

    assert _configure_pdf_styles()["AcademicBody"].alignment == TA_JUSTIFY
